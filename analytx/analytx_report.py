# -*- coding: utf-8 -*-
version = "18042026"

# %% imports

import locale

locale.setlocale(locale.LC_ALL, '')

try:
    import urllib.request
    import pandas as pd
    import os
    import re
    from io import BytesIO
    import openpyxl  # excell reading framework
    import docx
    from docx import Document  # document generation    row[0].paragraphs[0].runs[0].
    from docx.shared import Pt, Cm, Mm  # document generation
    from docx.enum.text import WD_ALIGN_PARAGRAPH  # document generation
    from docx.enum.table import WD_ALIGN_VERTICAL
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    from docx.shared import RGBColor
    from docx.oxml.shared import OxmlElement, qn
    import sys
    from sys import exit
except Exception as ex:
    print("ERROR: " + str(ex))
    print("  HELP: Local environment could require to install additional packages:")
    print("        command: pip install python-docx pandas openpyxl pandas.compat")
    print("  On pycharm, run the command from the lower panel console [>_] (venv) prompt")
    exit(1)

# GLOBAL VARS
document = Document() # global document we print on to generate output doc
debug = False # activate old style debug (print)
warns = [] # we add warns to be printed at the end of output doc

# Simple color palette for text
green_color = RGBColor(100, 150, 75)
gray_color = RGBColor(150, 150, 150)
red_color = RGBColor(200, 50, 50)
blue_color = RGBColor(79, 129, 189)
orange_color = RGBColor(250, 150, 50)


def cols_prepare(tab, cols):
    """
    Create a dictionary to register the columns we are using for later printing
    This is just for copy and paste for updating the dictionary in cols_validate
    It has no influence on program behaviour

    Usage:
    Set local var do_cols_prepare to True during test to generate it
    Test the application
    Print cols_prepare.idx_sheets
    Update the code accordingly
    Set local var do_cols_prepare to False

    :param tab: String with tab to add to the dictionary if it doesn't exist yet
    :param cols: List of cols to add to the dictionary if those don't exist yet
    :return: None
    """

    do_cols_prepare = False
    if not do_cols_prepare : return  # not needed until we re-generate the idx_sheets dictionary

    if not hasattr(cols_prepare, "idx_sheets"):
        cols_prepare.idx_sheets = {}  # it doesn't exist yet, so initialize it

    if tab in cols_prepare.idx_sheets.keys():
        idx_cols = cols_prepare.idx_sheets[tab]
    else:
        idx_cols = []

    if cols is not None:
        if isinstance(cols, list):
            for col in cols:
                print(str(tab) + ":" + str(col))
                if col not in idx_cols: idx_cols.append(col)
        else:
            print(str(tab) + ":" + str(cols))
            if cols not in idx_cols: idx_cols.append(cols)

        cols_prepare.idx_sheets[tab] = idx_cols

    return


required_sheets = {
#    Dictionary of required sheets and columns to be checked
#    Required by cols_validate and tabs_validate
#    Generated using cols_prepare during tests

    'vInfo': ['Memory', 'Datacenter', 'Cluster', 'Resource pool', 'NICs', 'VM',
              'OS according to the VMware Tools', 'Disks'],
    'vCPU': ['Hot Add', 'Hot Remove'],
    'vMemory': ['Hot Add', 'Ballooned'],
    'vMultiPath': ['Disk', 'Display name', 'Vendor', 'Model'],
    'vSource': ['Fullname', 'API version'],
    'vHost': ['ESX Version', '# CPU', 'Cores per CPU', '# Cores', 'CPU Model', 'Host',
              'Vendor', 'Model', 'Datacenter', 'Cluster', '# vCPUs', '# VMs total'],
    'vTools': ['Tools'],
    'vNIC': ['Datacenter', 'Driver', 'Speed', 'Duplex', 'MAC', 'Host'],
    'vNetwork': ['Network', 'VM'],
    'vDatastore': ['Type', 'Capacity MiB', 'Provisioned MiB', 'In Use MiB', 'Object ID'],
    'vDisk': ['Controller', 'VM', 'Capacity MiB'],
    'vPartition': ['Annotation', 'Capacity MiB', 'Disk', 'Free MiB'],
    'vHBA': ['Model', 'Type']
}


# def check_columns(file, df):

def cols_validate(tab, file, df):
    """
    Validate the required columns exist on spreadsheet sheet (tab)

    :param tab: the sheet (tab) name to be validated
    :param file: the original file (to enrich message)
    :param df: the dataframe to be validated
    :return: True if ok, False if a required column doesn't exist
    """
    data_ok = True
    if tab in required_sheets.keys():
        tab_cols = required_sheets[tab]
        for col in tab_cols:
            if not col in df.columns:
                print("ERROR: Column [" + col + "] is required in sheet [" + tab + "] in file [" + file + "]")
                data_ok = False
            else:
                df[col] = df[col].fillna(
                    'No' + col.replace(' ', '_'))  # replace empty value by "NoNetwork", "NoCluster", etc

    pcol = "none"
    for col in df.columns:
        # print(str(col)+",")
        if not isinstance(col, str):
            print("WARNING: Column [" + str(
            col) + "] just after [" + pcol + "] has wrong type in sheet [" + tab + "] in file [" + file + "] > Dropped")
            df.drop(columns=[col], inplace=True)
        pcol = col

    return data_ok


def tabs_validate(file, sheets):
    """
    Validate the required sheets (tabs) exist on spreadsheet

    :param file: the original file (to enrich message)
    :param sheets: a sheets dictionary
    :return: True if ok, False if a required sheet (tab) doesn't exist
    """

    data_ok = True
    for tab in required_sheets.keys():
        if tab not in sheets.keys():
            print("ERROR: Sheet [" + tab + "] is required in file [" + file + "]")
            data_ok = False
    return data_ok

#######################################################################
# %% PANDAS Functions
#

def try_search(ce, series, x, debug):
    """
    Utility function used in global_search leading with lists search and type errors
    :param ce: compiled regular expression
    :param series: the sheet name, this is only used for debug
    :param x: the element to search on
    :return:
    """

    result = None
    try:
        result = bool(ce.search(x))
    except TypeError:
        result = False

    if debug :
        if series == "Annotation": print("            try_search(ce," + series + "," + str(x) + ") returns " + str(result))

    return result


def global_search(sheets_dict, expression_dict, sheet_names_list=None):
    """
    Search each column in sheet_names_list sheets using expression_dict expressions where
    the key is a tag string representing the expression meaning, and the value is the expression string

    WARNINGS:

    - The query_expr doesn't support spaces or # symbol on column names so spaces should be provided as underscores and # as N.
    - "B" represents boundaries escape sequence (must be passed as double backlash followed by b)

    Expression syntax examples:

    - contains a string: 'string'
    - contains string1 or string2: 'string1|string2'
    - contains string1 and string2: '(?=.*string1)(?=.*string2)' #using non-consuming match
    - contains a word: 'BwordB' # using boundaries escape sequence. See warnings
    - contains word1 or word2: 'Bword1B|Bword2B' # using boundaries escape sequence. See warnings
    - contains word1 and word2: '(?=.*Bstring1B)(?=.*Bstring2B)' #using both non-consuming match and boundaries. See warnings

    :param sheets_dict: the global sheets dictionary
    :param expression_dict: a dictionary of "[expression-title][regular expression]" search terms
    :param sheet_names_list: a list of sheets (tabs) to search in. ALL if None
    :return: search results nested list
    """

    comb_search_results = []

    # Search in all sheets if sheet_names_list is None
    if sheet_names_list is None:
        sheet_names_list = []
        for name in sheets_dict.keys():
            s_name = name.split('@')[0]
            if s_name not in sheet_names_list: sheet_names_list.append(s_name)

    for tag, expression in expression_dict.items():
        ce = re.compile(expression, re.I)  # including the Ignorecase flag there is no need for upper() when using it
        term_search_results = []
        for sheet_name in sheet_names_list:
            df = combine_data_sheets(sheets_dict, sheet_name)
            # print("globalSearch(): " + sheet_name + "  " + expression + "==============================================================")
            for series_name, series in df.items():
                count = 0
                is_string = pd.api.types.is_string_dtype(df[series_name])
                is_list = pd.api.types.is_list_like(df[series_name])
                if is_string or is_list and str(df[series_name].dtype) == "object": #check is a searchable column
                    # print("globalSearch(): " + str(series_name) + " is_string=" + str(is_string) + " is_list=" + str(is_list))
                    result_df = pd.DataFrame([])
                    if is_string: result_df = df[df[series_name].str.contains(expression, na=True, case=False)]
                    if is_list: result_df = df[[try_search(ce, series_name, x, False) for x in df[series_name]]]

                    count = result_df.dropna(how='all').shape[0]

                    if count > 0: # add column results to current sheet results
                        # print("globalSearch() count=" + str(count) + " found in " + sheet_name +"/"+ series_name)
                        # print(result_df[series_name])
                        row = [tag + " (" + expression + ")", count, series_name, sheet_name]
                        term_search_results.append(row)
                        # print("         [" +expression+ "] " + str(count) + "[" + series_name + "/" + sheet_name + "]\n")
                # else: print("    " + str(series_name) + " is NOT String")
        comb_search_results = comb_search_results + term_search_results
        if len(term_search_results) == 0:
            row = [tag + " (" + expression + ")", 0, "", ""]
            comb_search_results.append(row)
        term_search_results = []

    return comb_search_results


def combine_data_sheets(sheets_dict, sheet, debug=False):
    """
    Populates a DataFrame which combines data from the sheets matching the sheet name.
    Aggregates multiple spreadsheets (multiple files) in one DataFrame.

    Args:
      sheets_dict: Dictionary containing all the Excel sheets.
      sheet: Name of the sheet to search for.

    Returns:
      A dataframe composed by the combination of all concatenated sheets
    """

    # create a local cache if it doesn't exist
    if not hasattr(combine_data_sheets, "cache"):
        combine_data_sheets.cache = {}  # cache dictionary doesn't exist yet, so initialize it

    # return cached pandas data frame from cache if it already exists
    if sheet in combine_data_sheets.cache.keys():
        if debug: print("combineDataSheets() retrieving " + sheet + " dataframe from cache")
        combined_df = combine_data_sheets.cache[sheet]
        return combined_df

    combined_df = pd.DataFrame()  # avoid 'in use before assignment' warning

    # Iterate through all the sheets in the dictionary
    for sheet_name, df in sheets_dict.items():
        # Check if the sheet name in dictionary ("sheetname@filename.xlsx") contains the searched name
        if sheet_name.startswith(sheet + "@") and not df.empty:
            if debug: print("    combineDataSheets() " + sheet + " in " + sheet_name)
            # Concatenate the content of this sheet into the combined DataFrame
            combined_df = pd.concat([combined_df, df])

    if debug: print("combineDataSheets() adding new " + sheet + " dataframe to cache")

    combine_data_sheets.cache[sheet] = combined_df

    # reindex after all
    combined_df.reset_index(inplace=True, drop=True)
    return combined_df


def calculate_percentage(df_dict, sheet, columns, ascending=False, debug=False):
    """
    Calculates the value and percentage represented by each unique combination of
    selected 'columns' in a 'sheet', ordered by percentage from highest to lowest.
    Args:
      df_dict: Dictionary containing all the Excel sheets.
      sheet: Name of the sheet to search for.
      columns: The columns (list) or column (string) name to analyze.
      ascending: None, False, True. How to sort based on Percentage (False by default)

    Returns:
      A list of tuples containing each unique value, count, and percentage, ordered by percentage.
      If columns parameter is a list, the returned 'value' will be a list, other case a string
    """
    cols_prepare(sheet, columns) # Ensure the columns presence in sheet is checked

    if debug: print(
        "calculate_percentage(df_dict," + sheet + "," + str(columns) + ", ascending=" + str(ascending) + ")")
    # Initialize an empty DataFrame to combine all the sheets that match the sheet_name_search
    combined_df = combine_data_sheets(df_dict, sheet)
    data = []

    # validate column...
    if isinstance(columns, list):
        for col in columns:
            if not col in combined_df.columns:
                print("ERROR calculate_percentage(): column [" + col + "] not found in sheet [" + sheet + "]")
                return data
            else:
                if (debug): print("column [" + col + "] found in sheet [" + sheet + "]")
                # Prevent crashing due to possible NaN
                combined_df[col] = combined_df[col].fillna('-')
    else:
        if not columns in combined_df.columns:
            print("ERROR calculate_percentage(): column [" + columns + "] not found in sheet [" + sheet + "]")
            return data
        else:
            if (debug): print("column [" + columns + "] found in sheet [" + sheet + "]")
            combined_df[columns] = combined_df[columns].fillna('-')

    # Calculate value counts and percentages
    counts = combined_df[columns].value_counts(dropna=True)  # Not sure if we are loosing some info with dropna=True
    counts = counts.where(pd.notnull(counts), 'Unknown')  # replace nan by none
    total_count = counts.sum()
    percentages = (counts / total_count) * 100

    # Create a list of tuples with value, count, and percentage
    for value, count in counts.items():
        percentage = percentages[value]
        data.append((value, count, percentage))

    # Sort by percentage in descending order
    if ascending == True: data.sort(key=lambda x: x[2])
    if ascending == False: data.sort(key=lambda x: x[2], reverse=True)

    return data


def get_rows(df_dict, sheet, key_columns=None, columns=None, query_expr=None, contains_column=None, contains_expr=None,
             contains_case=True, debug=False):
    """
    Function to extract rows, based on a query with an option to extract only unique elements on a key.

    WARNINGS:

    - The query_expr doesn't support spaces or # symbol on column names so spaces should be provided as underscores and # as N.
    - "B" represents boundaries escape sequence (must be passed as double backlash followed by b)

    Contains examples:

    - contains a string: 'string'
    - contains string1 or string2: 'string1|string2'
    - contains string1 and string2: '(?=.*string1)(?=.*string2)' #using non-consuming match
    - contains a word: 'BwordB' # using boundaries escape sequence. See warning
    - contains word1 or word2: 'Bword1B|Bword2B' # using boundaries escape sequence. See warning
    - contains word1 and word2: '(?=.*Bstring1B)(?=.*Bstring2B)' #using both non-consuming match and boundaries. See warning

    Updating data is only possible when "columns" == None because in other case it will return a copy
    for index, row in df.iterrows():  #iterate
        df.at[index,'Column name'] = new_value   #update

    :param df_dict: Dictionary containing all the Excel sheets.
    :param sheet: Name of the sheet to search for and get rows.
    :param key_columns: List of Key columns to remove duplicates on its combination.
    :param columns: List of columns to return (default is all).
    :param query_expr: Query expression for DataFrame.query
    :param contains_column: Column to be searched with contains_expr
    :param contains_expr: Expression to search for in the specified column. OR "term1|term2", AND '(?=.*term1)(?=.*term2|.*term3)'
    :param contains_case: Boolean to ignore case on contains expression (only)
    :param debug: Boolean to print debug information.
    :return: resulting df

    """
    cols_prepare(sheet, key_columns)
    cols_prepare(sheet, columns)

    if debug:
        print("get_rows(df_dict, sheet_name_search:" + str(sheet) + ", key_column:" + str(
            key_columns) + ", query_expr:" + str(
            query_expr) + ", contains_column:" + str(contains_column) + ", contains_expr:" + str(
            contains_expr) + ", case:" + str(contains_case) + ")")

    # Initialize an empty DataFrame to combine all the sheets that match the sheet_name_search
    combined_df = combine_data_sheets(df_dict, sheet)
    if combined_df.size == 0: return combined_df

    if debug:
        print("get_rows():df.SIZE:" + str(combined_df.size))
        # print(combined_df)

    if query_expr is not None:
        combined_df = combined_df.query(query_expr)

    if contains_expr is not None:
        if debug: print("get_rows(): Contains: Column:" + str(contains_column) + "Expression:" + str(contains_expr))
        combined_df = combined_df[combined_df[contains_column].str.contains(contains_expr, na=False, case=contains_case)]

    if debug:
        print("get_rows():AFTER FILTER")
        # print(combined_df)

    # Remove duplicates based on the key_columns
    if key_columns is not None:
        if isinstance(key_columns, list):
            combined_df = combined_df.drop_duplicates(subset=key_columns)
        else:
            combined_df = combined_df.drop_duplicates(subset=[key_columns])

    if columns is not None:
        if debug: print("get_rows() columns=" + str(columns))
        combined_df = combined_df[columns]

    if debug:
        print("get_rows():AFTER UNIQUE")
        # print(combined_unique_df)

    return combined_df


def count_rows(df_dict, sheet, key_columns=None, count_unique=True, query_expr=None, contains_column=None,
               contains_expr=None, contains_case=True, debug=False):
    """
    Function to count rows, based on a query with an option to extract only unique elements on a key.

    :param df_dict: Dictionary containing all the Excel sheets.
    :param sheet: Name of the sheet to search for and count rows.
    :param key_columns: Key columns combination to remove duplicates on.
    :param count_unique: Boolean to count only unique rows (True) or all rows (False).
    :param query_expr: Expresion for DataFrame.query
    :param contains_column: Column to be searched with contains_expr
    :param contains_expr: Expression to search for in the specified column. OR "term1|term2", AND '(?=.*term1)(?=.*term2|.*term3)'
    :param contains_case: Boolean to ignore case on contains expression (only)
    :param debug: Boolean to print debug information.
    :return: Number of non-empty rows (unique or all depending on the flag) in the specified sheet(s).
    """
    if debug:
        print("count_rows(df_dict," + sheet + "," + key_columns + "," + str(count_unique) + "," + str(query_expr))

    combined_unique_df = get_rows(df_dict, sheet, key_columns, None, query_expr, contains_column, contains_expr, contains_case,
                                  False)

    if debug:
        print(">>>>>>>count_rows() combined_unique_df")
        if contains_column is not None:
            showcols = []
            showcols.append(key_columns)
            showcols.append(contains_column)
            # print(str(combined_unique_df[showcols]))
        # else:
        #  print(str(combined_unique_df))
        print("<<<<<<<count_rows() combined_unique_df")

    # Count non-empty rows in the combined DataFrame
    total_rows = combined_unique_df.dropna(how='all').shape[0]

    # Display the result if debug is activated
    if debug:
        print(
            f"count_rows(): Total {'unique' if count_unique else 'all'} non-empty rows in the sheet '{sheet}': {total_rows}")

    return total_rows


def sum_rows(df_dict, sheet_name, key_column, debug=False):
    """
    Calculates the sum of key_colum values in sheet_name sheet
    :param df_dict:  Global sheets dictionary
    :param sheet_name: the sheet (tab) name
    :param key_column: the column to sum values
    :param debug:

    :return: sum result
    """

    cols_prepare(sheet_name, key_column)

    count = 0
    columnExists = False
    sheetExists = False

    combined_df = combine_data_sheets(df_dict, sheet_name)
    count = combined_df[key_column].sum()

    if debug: print("sum_rows('" + sheet_name + "','" + key_column + "') =" + str(count))
    return count


def groupby(df_dict, sheet='vInfo', columns=None, trunk=True, ascending=None, sum=False, result_name=None, debug=False):
    """
    group by a list of columns in a sheet

    ie: groupby(myDict, 'vInfo', ["Datacenter","Cluster","Host","VM"],True,True)

    The last column will show the sum or count of instances on the previous column.
    In the example, "VM" would change to "Count" in the returning df

    Args:
      df_dict: Dictionary containing all the Excel sheets.
      sheet: Name of the sheet to search for.
      columns: Names of the columns to group by
      trunk: Return only the columns included in "columns" otherwise return all
      ascending: Sort result by last column Count ascending (True), descending(False) or not sort at all (None)
      sum: Defines if this is required to count (False) or sum (True) the last column values
      result_name: Optional name to be used for the sum or count column. "Sum" or "Count" if None

    Returns:
      A grouped df
    """
    cols_prepare(sheet, columns) #to validate columns we use

    if debug: print("groupby(df_dict," + sheet + "," + str(columns) + ")")

    # Initialize an empty DataFrame to combine all the sheets that match the sheet_name_search
    combined_df = combine_data_sheets(df_dict, sheet)

    lastcol = columns[-1]

    new_columns = columns.copy()
    new_columns.pop()

    if sum:
        grouped = combined_df.groupby(new_columns)[lastcol].sum().reset_index()
        if result_name is None: result_name = 'Sum'
    else:
        grouped = combined_df.groupby(new_columns).count().reset_index()
        if result_name is None: result_name = 'Count'

    # if trunk return the requested columns only
    if trunk: grouped = grouped[columns]

    if result_name is not None and len(result_name) > 0:
        grouped.rename(columns={lastcol: result_name}, inplace=True)
        # print("groupby(df_dict," + sheet + "," + str(columns) + ") SORT by:"+result_name)
        if ascending is True: grouped = grouped.sort_values(by=[result_name], ascending=True)
        if ascending is False: grouped = grouped.sort_values(by=[result_name], ascending=False)

    if debug: print("\n" + str(grouped) + "\n\n")

    return grouped


def groupby_df(combined_df, columns=None, trunk=True, ascending=None, sum=False, result_name=None, debug=False):
    """
    group by a list of columns in a sheet
    example: groupby(myDict, 'vmInfo', ["Datacenter","Cluster","Host","VM"])
    The last column will be shown as the sum of instances on the previous column.
    In the example, "VM" would show how much VMs are on each "Host"

    Args:
      combined_df: DataFrame to group
      columns: Names of the columns to group by
      trunk: Return only the columns included in "columns" otherwise return all
      ascending: Sort result by last column Count ascending (True), descending(False) or not sort at all (None)
      sum: Defines if this is required to count (False) or sum (True) the last column values

    Returns:
      A list of tuples containing the OS value, count, and percentage, ordered by percentage.
    """

    if debug: print("groupby_df(combined_df," + str(columns) + ")")

    lastcol = columns[-1]
    new_columns = columns.copy()
    new_columns.pop()

    # print("groupby() combined")
    # print(combined_df)
    if sum:
        grouped = combined_df.groupby(new_columns)[lastcol].sum().reset_index()
        if result_name is None: result_name = 'Sum'
    else:
        grouped = combined_df.groupby(new_columns).count().reset_index()
        if result_name is None: result_name = 'Count'

    if trunk: grouped = grouped[columns]

    # Rename last column as Count
    old_name = grouped.columns[len(columns) - 1]
    if result_name is not None and len(result_name) > 0:
        grouped.rename(columns={old_name: result_name}, inplace=True)
        if ascending is True: grouped = grouped.sort_values(by=[result_name], ascending=True)
        if ascending is False: grouped = grouped.sort_values(by=[result_name], ascending=False)

    if debug: print("\n" + str(grouped) + "\n\n")

    return grouped


'''
############################################################################################
# %% UTILITY Functions
############################################################################################
'''


def lapse(msg="", on=None):
    """
    Print message and time lapse from previous lapse invocation.

    - To activate lapse printing: lapse(on=True)
    - To deactivate lapse printing: lapse(on=False)
    - To print a lapse: lapse("Lapse message")

    :param msg: lapse message
    :param on: True/False to activate/deactivate lapse
    :return: None
    """
    from time import time

    if not hasattr(lapse, "on"):
        lapse.on = False

    if on is not None: lapse.on = on

    if lapse.on is False:
        return

    if not hasattr(lapse, "last_time"):
        lapse.last_time = time()  # doesn't exist yet, so initialize it
        print("lapse () Initialized")
    else:
        msg = 'elapsed (%f)' % (time() - lapse.last_time) + " " + msg
        print(msg)
        lapse.last_time = time()

'''
############################################################################################
# %% Hyperlinks related Functions
############################################################################################
'''

def add_hyperlink(paragraph, text, url):
    """
    Add url link to paragraph text

    :param paragraph: The text paragraph
    :param text: The text
    :param url: The url
    :return: The new run added to paragraph
    """
    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

    # Create a new run object (a wrapper over a 'w:r' element)
    new_run = docx.text.run.Run(
        docx.oxml.shared.OxmlElement('w:r'), paragraph)
    new_run.text = text
    # new_run.font.underline = True
    # new_run.font.color.rgb = blue_color

    # Set the run's style to the builtin hyperlink style, defining it if necessary
    #       new_run.style = get_or_create_hyperlink_style(part.document)
    # Alternatively, set the run's formatting explicitly
    # new_run.font.color.rgb = docx.shared.RGBColor(0, 0, 255)
    # new_run.font.underline = True

    # Join all the xml elements together
    hyperlink.append(new_run._element)
    paragraph._p.append(hyperlink)
    return new_run


# This is only needed if you're using the builtin style above
def get_or_create_hyperlink_style(d):
    """
    If this document had no hyperlinks so far, the builtin
    Hyperlink style will likely be missing and we need to add it.
    There's no predefined value, different Word versions
    define it differently.
    This version is how Word 2019 defines it in the
    default theme, excluding a theme reference.
    """
    if "Hyperlink" not in d.styles:
        if "Default Character Font" not in d.styles:
            ds = d.styles.add_style("Default Character Font",
                                    docx.enum.style.WD_STYLE_TYPE.CHARACTER,
                                    True)
            ds.element.set(docx.oxml.shared.qn('w:default'), "1")
            ds.priority = 1
            ds.hidden = True
            ds.unhide_when_used = True
            del ds
        hs = d.styles.add_style("Hyperlink",
                                docx.enum.style.WD_STYLE_TYPE.CHARACTER,
                                True)
        hs.base_style = d.styles["Default Character Font"]
        hs.unhide_when_used = True
        hs.font.color.rgb = docx.shared.RGBColor(0x05, 0x63, 0xC1)
        hs.font.underline = True
        del hs

    return "Hyperlink"


def urlRead(base, path):
    """
    Read url contents, EXPERIMENTAL !!

    usage: urlRead("https://catalog.redhat.com/search?gs&q=","PowerEdge R1750")

    :param base: url base without escape parsing requirements
    :param path: url part to be escaped, usually the variable part

    :return: a file with results
    """
    # urlRead("https://catalog.redhat.com/search?gs&q=","PowerEdge R1750")
    link = base + urllib.parse.quote(path, safe='/', encoding=None, errors=None)
    # print("base=["+base+"] + path=["+path+"] link=["+link+"]")
    f = urllib.request.urlopen(link)

    myfile = str(f.read())
    res = myfile.find("No results found")
    # print("RESULT:"+str(res))

    return myfile


"""
####################################################################################
# %% REPORT (Docx) Functions
####################################################################################
"""

def keep_table_on_one_page(doc):
    tags = doc.element.xpath('//w:tr[position() < last()]/w:tc/w:p')
    for tag in tags:
        ppr = tag.get_or_add_pPr()
        ppr.keepNext_val = True


def set_cell_background(cell=None, color="CCCCCC"):
    from docx.oxml.ns import nsdecls
    from docx.oxml import parse_xml

    tblCell = cell._tc
    tblCellProperties = tblCell.get_or_add_tcPr()
    clShading = OxmlElement('w:shd')
    clShading.set(qn('w:fill'), color)  # Hex of Dark Blue Shade {R:0x00, G:0x51, B:0x9E}
    tblCellProperties.append(clShading)
    return cell

    # shading_elm = parse_xml(r'<w:shd {} w:fill="cccccc"/>'.format(nsdecls('w')))
    # cell._tc.get_or_add_tcPr().append(shading_elm)


def set_row_height(row, height=1000):
    # https://stackoverflow.com/questions/37532283/python-how-to-adjust-row-height-of-table-in-docx
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    trHeight = OxmlElement('w:trHeight')
    trHeight.set(qn('w:val'), str(height))
    trHeight.set(qn('w:hRule'), "atLeast")
    trPr.append(trHeight)


def set_table_cells_margins(table, **kwargs):
    for row in table.rows:  ###
        for cell in row.cells:
            set_cell_margins(cell, **kwargs)


def set_cell_margins(cell, **kwargs):
    """
    set_cell_margins(cell, top=50, start=50, bottom=50, end=50)

    provided values are in twentieths of a point (1/1440 of an inch).
    read more here: http://officeopenxml.com/WPtableCellMargins.php

    Args:

    cell:  actual cell instance you want to modify

    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')

    for m in [
        "top",
        "start",
        "bottom",
        "end",
    ]:
        if m in kwargs:
            node = OxmlElement("w:{}".format(m))
            node.set(qn('w:w'), str(kwargs.get(m)))
            node.set(qn('w:type'), 'dxa')
            tcMar.append(node)

    tcPr.append(tcMar)


def set_table_borders(table, **kwargs):
    """
    Set table`s borders

    set_table_borders(
        table,
        top={"sz": 12, "val": "single", "color": "#FF0000", "space": "0"},
        bottom={"sz": 12, "color": "#00FF00", "val": "single"},
        start={"sz": 24, "val": "dashed", "shadow": "true"},
        end={"sz": 12, "val": "dashed"},
    )
    """
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(cell, **kwargs)


def set_run_title_style(run):
    run.font.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = blue_color
    return run


def set_table_title(table, title):
    cell = None
    if title is not None:
        header_row = table.add_row()
        cell = header_row.cells[0]
        cell.text = title
        set_run_title_style(cell.paragraphs[0].runs[0])
        # cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.merge(header_row.cells[-1])
        insertion_row = table.rows[0]._tr
        insertion_row.addprevious(table.rows[-1]._tr)
        set_cell_border(
            cell,
            top={"sz": 0, "color": "#CCCCCC", "val": "single"},
            bottom={"sz": 10, "color": "#CCCCCC", "val": "single"},
            start={"sz": 0, "color": "#CCCCCC", "val": "single"},
            end={"sz": 0, "color": "#CCCCCC", "val": "single"}
        )
    return cell


def set_cell_border(cell, **kwargs):
    """
    Set cell`s border
    Usage:

    set_cell_border(
        cell,
        top={"sz": 12, "val": "single", "color": "#FF0000", "space": "0"},
        bottom={"sz": 12, "color": "#00FF00", "val": "single"},
        start={"sz": 24, "val": "dashed", "shadow": "true"},
        end={"sz": 12, "val": "dashed"},
    )
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    # check for tag existnace, if none found, then create one
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)

    # list over all available tags
    for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)

            # check for tag existnace, if none found, then create one
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)

            # looks like order of attributes is important
            for key in ["sz", "val", "color", "space", "shadow"]:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))


def table_from_df(df, document, grid=False, columnWidths=None, title=None):
    """
    Populates a document table from pandas DataFrame data

    :param df: the dataframe to be printed
    :param document: the document to print in
    :param grid: do we want to print with grid style or not
    :param columnWidths: list with desired with for each column (None=Automatic)
    :param title: the resulting table title (or None)
    :return: the document added table object
    """
    rows = df.shape[0] + 1
    cols = df.shape[1]

    table = document.add_table(rows, cols)
    if grid: table.style = 'Table Grid'

    r = 0
    for row in table.rows:
        c = 0
        for cell in row.cells:

            if r == 0:  # Headers
                text = df.columns[c]
                cell.text = str(text)
                cell.paragraphs[0].runs[0].font.bold = True
                if columnWidths is not None: table.columns[c].width = Pt(columnWidths[c])
            else:  # Data
                text = df.iloc[[r - 1]].values[0][c]
                # print("TYPE of CELL is "+str(type(text)))
                cell.text = str(text)
                if isinstance(text, int) or isinstance(text, float): cell.paragraphs[
                    0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

            if columnWidths is not None: cell.width = Pt(columnWidths[c])

            set_cell_border(
                cell,
                top={"sz": 10, "color": "#CCCCCC", "val": "single"},
                bottom={"sz": 10, "color": "#CCCCCC", "val": "single"},
                start={"sz": 10, "color": "#CCCCCC", "val": "single"},
                end={"sz": 10, "color": "#CCCCCC", "val": "single"}
            )
            c = c + 1
        r = r + 1

    set_table_title(table, title)

    return table


def add_percentage_table(document, mylist, data, rowColor=None, style=None, borderColor=None, links=None, title=None):
    """
    Populates a document table from a list of "value, count, percentage" rows
    :param document: the document to add the table to
    :param mylist: the input data list. Each row contains value, count and percentage
    :param data:
    :param rowColor:
    :param style:
    :param borderColor:
    :param links:
    :param title:
    :return:
    """
    table = document.add_table(1, 3)
    table.style = style
    table.autofit = True
    table.allow_autofit = True

    w0 = Pt(400)
    w1 = Pt(50)
    w2 = Pt(50)

    # w0=Inches(3.5)
    # w1=Inches(1)
    # w2=Inches(1)

    # Open Office needs it
    table.columns[0].width = w0
    table.columns[1].width = w1
    table.columns[2].width = w2

    row = table.rows[0].cells
    row[0].text = data
    row[0].width = w0
    row[0].paragraphs[0].runs[0].font.bold = True
    row[1].text = 'Count'
    row[1].width = w1
    row[1].paragraphs[0].runs[0].font.bold = True
    row[2].text = '%'
    row[2].width = w2
    row[2].paragraphs[0].runs[0].font.bold = True

    # TODO: Validate MS Word behaviour, this works ok with GDocs and Libre Office but not tested with MS Office
    table.cell(0, 0).width = w0
    table.cell(0, 1).width = w1
    table.cell(0, 2).width = w2

    i = 0
    for value, count, percentage in mylist:
        row = table.add_row().cells

        # avoid data concatenation without separator
        if isinstance(value, tuple) and len(value) > 1:
            vText = ''
            for val in value: vText = vText + str(val).strip() + " / "
            vText = vText[0:len(vText) - 3]  # clean last /
        else:
            vText = str(value)

        if links is None:
            row[0].text = vText
        else:
            paragraph = row[0].paragraphs[0]
            run = add_hyperlink(paragraph, vText, links[i])
            run.font.underline = True
            run.font.color.rgb = blue_color

        if rowColor is not None:
            row[0].paragraphs[0].runs[0].font.color.rgb = rowColor[i]

        row[1].text = str(count)
        row[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        row[2].text = "{pct:.2f}".format(pct=percentage) + "%"
        row[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        # row[0].width = w0
        # row[1].width = w1
        # row[2].width = w2

        i = i + 1

    if borderColor is None: borderColor = "#CCCCCC"

    set_table_borders(table,
                      top={"sz": 10, "color": borderColor, "val": "single"},
                      bottom={"sz": 10, "color": borderColor, "val": "single"},
                      start={"sz": 10, "color": borderColor, "val": "single"},
                      end={"sz": 10, "color": borderColor, "val": "single"}
                      )

    set_table_title(table, title)

    return table


def table_from_list(list, document, grid=False, columns_width=None, columns_wd_align=None, title=None):
    """
      Populates a document table from a list of rows

      Args:
          list: A nested list of rows and columns including headers
          document: The document we are going to write to
          grid: Set style as "Table Grid"
          columns_width: A list with width in points for each column
          columns_wd_align: A list with WD_ALIGN_PARAGRAPH.xxx for each column

      Returns:
          The created table which is writable
    """

    rows = len(list)
    cols = len(list[0])

    # print("table_from_list() rows="+str(rows))
    # print("table_from_list() cols="+str(cols))
    # print(list)

    table = document.add_table(rows, cols)
    if grid: table.style = 'Table Grid'

    r = 0
    for row in table.rows:
        c = 0
        for cell in row.cells:
            # print("["+str(r)+","+str(c)+"] ", end=None)
            cell.text = str(list[r][c])

            if columns_width is not None:
                if r == 0:
                    table.columns[c].width = Pt(columns_width[c])  # for Libre Office
                    cell.paragraphs[0].runs[0].font.bold = True
                else:
                    if columns_wd_align != None and columns_wd_align[c] != None:
                        cell.paragraphs[0].alignment = columns_wd_align[c]
                cell.width = Pt(columns_width[c])

            set_cell_border(
                cell,
                top={"sz": 10, "color": "#CCCCCC", "val": "single"},
                bottom={"sz": 10, "color": "#CCCCCC", "val": "single"},
                start={"sz": 10, "color": "#CCCCCC", "val": "single"},
                end={"sz": 10, "color": "#CCCCCC", "val": "single"}
            )
            c = c + 1
        r = r + 1

    set_table_title(table, title)

    return table


def table_from_dict(dict, document, grid=False, columns_width=None, keyHeader="Key", dataHeader="Data", align_r=None,
                    title=None):
    """
    Populates a 'key: value style' document table from a dictionary
    Args:
        dict: A nested list of rows and columns including headers
        document: The document we are going to write to
        grid: Set style as "Table Grid"
        columns_width: A list with width in points for each column
        keyHeader: The header text for key column
        dataHeader: The header text for data column
        align_r: True is to right align data column, None or False for left align
    Returns:
        The created table which is writable
    """

    rows = len(dict)

    # print("table_from_list() rows="+str(rows))
    # print("table_from_list() cols="+str(cols))
    # print(list)

    table = document.add_table(rows + 1, 2)
    if grid: table.style = 'Table Grid'

    myKeys = list(dict.keys())
    myData = list(dict.values())

    # print(myKeys)
    # print(myData)

    r = 0
    for row in table.rows:
        # print("ROW "+str(r)+" of "+str(rows) + " myKeys["+str(r)+"-1]='"+myKeys[r-1]+"'")
        if r == 0:
            row.cells[0].text = keyHeader
            row.cells[1].text = dataHeader
            row.cells[0].paragraphs[0].runs[0].font.bold = True
            row.cells[1].paragraphs[0].runs[0].font.bold = True

        else:
            row.cells[0].text = str(myKeys[r - 1])
            row.cells[1].text = str(myData[r - 1])
            if align_r == True: row.cells[1].alignment = WD_ALIGN_PARAGRAPH.RIGHT

        if columns_width is not None:
            row.cells[0].width = Pt(columns_width[0])
            row.cells[1].width = Pt(columns_width[1])

        set_cell_border(
            row.cells[0],
            top={"sz": 10, "color": "#CCCCCC", "val": "single"},
            bottom={"sz": 10, "color": "#CCCCCC", "val": "single"},
            start={"sz": 10, "color": "#CCCCCC", "val": "single"},
            end={"sz": 10, "color": "#CCCCCC", "val": "single"}
        )

        set_cell_border(
            row.cells[1],
            top={"sz": 10, "color": "#CCCCCC", "val": "single"},
            bottom={"sz": 10, "color": "#CCCCCC", "val": "single"},
            start={"sz": 10, "color": "#CCCCCC", "val": "single"},
            end={"sz": 10, "color": "#CCCCCC", "val": "single"}
        )
        r = r + 1

    set_table_title(table, title)

    return table


def stops_table(list, document, grid=False, stops=None, columns_align_r=None):
    # TODO unfinished function
    """
    Generate shortTable grouping segments based in tabs-stop values. <= x, >x & <x1, >=x1 & < x2, ...
    EXPERIMENTAL !!!

    Args:
        list: A nested list of rows and columns including headers
        document: The document we are going to write to
        grid: Set style as "Table Grid"
        columns_width: A list with width in points for each column
        columns_align_r: A list with WD_ALIGN_PARAGRAPH.xxx for each column

    Returns:
        table: The created table which is writable

  See: https://stackoverflow.com/questions/67861989/apply-comparison-operator-from-string
       https://stackoverflow.com/questions/66950993/input-a-string-with-a-comparison-operator-and-perform-the-comparison-on-an-array

  To be used from: "[vNetwork] Network/VM (VMs per Network)",

  for row in network_VM.itertuples():
    if (row.Count < 5):
      index = "just "+str(row.Count)
      currentValue = network_VM_shortList.get(index,0)
      newValue = currentValue + row.Network
      network_VM_shortList.update({index:newValue})
    if row.Count >= 5 and row.Count < 10:
      index = ">=5..<10"
      ...
    if row.Count > 100:
      ...

    """

    print("stopsTable() ERROR: NOT IMPLEMENTED")
    return None


def add_h(txt='', p=None):
    """
    Add a header (txt) to global document
    :param txt: The header text
    :param p: The paragraph to add the header to. A new one is created if None is passed in


    :return: header paragraph
    """
    space_before = 6
    space_after = 6


    if p is None: p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(space_before)  # Espacio antes del párrafo
    p.add_run(txt).bold = True
    p.paragraph_format.space_after = Pt(space_after)  # Espacio detrás del párrafo
    return p


def add_p():
    """
    Add a paragraph to global document
    :return: paragraph
    """
    space_after = 6

    # Agregar un párrafo
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(space_after)  # Espacio detrás del párrafo
    return p


'''
########################################################################################################
# INITIALIZE FUNCTIONS
########################################################################################################
'''


def _ingest_excel_workbook(sheets, basename, sheets_dict):
    """
    Merge one workbook's sheets into sheets_dict (mutates sheets_dict).
    Returns (data_ok, number_of_sheets_stored).
    """
    data_ok = True
    if not tabs_validate(basename, sheets):
        data_ok = False

    sheet_count = 0
    for sheet_name, df in sheets.items():
        sheet_key = f"{sheet_name}@{basename}"
        print("        storing sheet [" + sheet_key + "] in dictionary")

        if not cols_validate(sheet_name, basename, df):
            data_ok = False

        df.columns = [c.replace(' ', '_') for c in df.columns]
        df.columns = [c.replace('#', 'N') for c in df.columns]

        sheets_dict[sheet_key] = df
        sheet_count += 1

    return data_ok, sheet_count


def load_spreadsheets(path):
    '''
    #######################################################################
    # %% LOAD RVTools spreadsheets
    #
    '''

    sheets_dict = {}  # Spreadsheets dictionary

    # obtain the list of files
    files = [f for f in os.listdir(path) if f.endswith('.xlsx') or f.endswith('.xlsm')]

    sheetCount = 0
    filesCount = 0
    data_ok = True
    for file in files:
        file_path = os.path.join(path, file)
        print("    loading [" + file_path + "]")

        sheets = pd.read_excel(file_path, sheet_name=None, engine='openpyxl')
        ok, added = _ingest_excel_workbook(sheets, os.path.basename(file), sheets_dict)
        if not ok:
            data_ok = False
        sheetCount += added
        filesCount = filesCount + 1

    if filesCount == 0:
        print("No .xlsx or .xlsm files found in [" + path + "]")
        exit(2)
    else:
        print(str(sheetCount) + " sheets loaded from " + str(filesCount) + " files")
        if not data_ok:
            print("\nAborting due to incomplete data. \nPlease, review previous errors, fix it and try again")
            exit(1)

    return sheets_dict


def load_spreadsheets_bytes(content, original_filename, lang="en"):
    """
    Load a single RVTools .xlsx/.xlsm from bytes (e.g. web upload).
    Raises ValueError if the file is missing, unreadable, or fails validation.
    lang: 'en' or 'de' for user-facing error messages.
    """
    L = "de" if (lang or "").lower().startswith("de") else "en"
    name = os.path.basename(original_filename or "") or "rvtools.xlsx"
    lower = name.lower()
    if not (lower.endswith('.xlsx') or lower.endswith('.xlsm')):
        if L == "de":
            raise ValueError("Bitte eine RVTools-Datei im Format .xlsx oder .xlsm hochladen.")
        raise ValueError("Please upload an RVTools export file in .xlsx or .xlsm format.")

    sheets_dict = {}
    try:
        bio = BytesIO(content)
        sheets = pd.read_excel(bio, sheet_name=None, engine='openpyxl')
    except Exception as ex:
        if L == "de":
            raise ValueError("Die Excel-Datei konnte nicht gelesen werden: " + str(ex)) from ex
        raise ValueError("The Excel file could not be read: " + str(ex)) from ex

    ok, _ = _ingest_excel_workbook(sheets, name, sheets_dict)
    if not ok:
        if L == "de":
            raise ValueError(
                "Die Datei entspricht nicht dem erwarteten RVTools-Format (fehlende Tabellen oder Spalten). "
                "Details stehen in der Server-Konsole."
            )
        raise ValueError(
            "The file does not match the expected RVTools layout (missing sheets or columns). "
            "See the server console for details."
        )
    return sheets_dict


def clean_and_fix_data(sheets_dict):
    """
    Clean some column values for later convenience

    """
    # remove everything between '(' and ')' in storage Display Name
    vMultiPath_df = combine_data_sheets(sheets_dict, 'vMultiPath')
    if 'Display_name' not in vMultiPath_df.columns:
        vMultiPath_df['Display_name'] = '-'
    vMultiPath_df['Display_name'] = vMultiPath_df['Display_name'].str.replace(r'\(.*\)', '', regex=True)

    # remove "VMware Virtual Processor" when exist (only few cases)
    vHost_df = combine_data_sheets(sheets_dict, 'vHost')
    vHost_df.drop(vHost_df[vHost_df["CPU_Model"].str.contains("VMware")].index, axis=0, inplace=True)

    # Fix empty values. Empty columns will be deleted by "combine_data_sheets()"
    for sheet in ["vInfo", "vCPU", "vHost", "vMemory", "vDisk", "vPartition", "vNetwork", "vHBA", "vNIC", "vMultiPath"]:
        df = combine_data_sheets(sheets_dict, sheet)
        df.replace({'Cluster': {None: 'NoCluster'}}, inplace=True)
        # df.replace({'Cluster': {'': 'NoCluster'}}, inplace=True)

    for sheet in ["vInfo", "vCPU", "vMemory", "vDisk", "vPartition", "vNetwork"]:
        df = combine_data_sheets(sheets_dict, sheet)
        if 'Annotation' not in df.columns:
            continue
        df['Annotation'] = df['Annotation'].fillna('').astype(str).str.replace('.*@.*.com|.*@.*.es', 'xxx@acme.com', regex=True)

    vHBA_df = combine_data_sheets(sheets_dict, 'vHBA')
    vHBA_df.Type = vHBA_df.Type.fillna('-')

    #######################################################
    # Include Display_name column in vDatastore based on Type and Address
    #   Type == NFS  then Display_name is the server part of vDatastore.Address
    #   Type == VMFS then Display_name is copied from vMultipath.Display_name where vMultipath.Disk == vDatastore.Address
    #   Type == VSAN then Display_name is the vDatastore.Name
    vDatastore_df = combine_data_sheets(sheets_dict, 'vDatastore')
    values = []
    dname = "UNKNOWN"

    # Iterate Datastore Address
    for index, row in vDatastore_df.iterrows():
        vtype = str(row['Type']).strip()
        addr = str(row['Address']).strip()
        acc = str(row['Accessible'])
        name = str(row['Name'])
        dname = ''

        if vtype.upper() == "NFS":
            dname = re.sub(" /.*", "", addr)
        elif vtype.upper() == "VMFS":
            vMultiPaths = get_rows(sheets_dict, sheet='vMultiPath',
                                   key_columns='Disk', columns=['Host', 'Disk', 'Display_name', 'Vendor', 'Model'],
                                   query_expr="Disk=='" + addr + "'")
            for multipath in vMultiPaths.itertuples():
                dname = multipath.Display_name
        elif vtype.upper() == "VSAN":
            dname = str(row['Name']).strip()
        else:
            dname = "UNKNOWN TYPE  [" + vtype + "]"

        if str(dname).strip().startswith('NoDisplay'):  # SEE cols_validate() about 'NoDisplay' value
            warns.append(
                "Datastore " + name + " Type:" + vtype.upper() + " with addr [" + addr + "] Not Accessible or Disconnected")
            dname = "Not Accessible or Disconnected"  # Probably Accessible == False or "# Hosts" = 0

        values.append(dname)

    vDatastore_df.insert(3, "Display_name", values, True)

    return


def anonymize_names(sheets_dict, anonymize=False):
    """
    EXPERIMENTAL !!
    Anonymize data like Datacenter, cluster and VMs names
    :param sheets_dict:
    :param anonymize:
    :return:
    """
    #######################################################
    # Anonymize
    #######################################################

    if not anonymize: return None

    data = {}

    anonym_datacenter = {}
    anonym_cluster = {}
    seqDatacenter = 0
    seqCluster = 0

    #######################################################
    if debug: print("Anonymize vCluster: Name")
    #######################################################
    vCluster_df = combine_data_sheets(sheets_dict, 'vCluster')

    # Iterate vInfo for Datacenter and Cluster
    for index, row in vCluster_df.iterrows():
        vCluster = str(row['Name']).strip()
        if anonym_cluster.get(vCluster) is None:
            seqCluster = seqCluster + 1
            anonym_cluster[vCluster] = "Cluster-" + '{n:03}'.format(n=seqCluster)
            # print("Anonym: "+vDatacenter+"/"+vCluster+" as "+anonym_datacenter[vDatacenter]+"/"+anonym_cluster[vCluster])
        vCluster_df.at[index, "Name"] = anonym_cluster[vCluster]

    ########################################################
    if debug: print("Anonymize vInfo: Datacenter Cluster")  #
    ########################################################
    vInfo_df = combine_data_sheets(sheets_dict, 'vInfo')

    # Iterate vInfo for Datacenter and Cluster
    for index, row in vInfo_df.iterrows():

        vDatacenter = str(row['Datacenter']).strip()
        # if Datacenter not in anonymization dictionary add it
        if anonym_datacenter.get(vDatacenter) is None:
            seqDatacenter = seqDatacenter + 1
            anonym_datacenter[vDatacenter] = "Datacenter-" + '{n:02}'.format(n=seqDatacenter)
            # print("Anonym: " + vDatacenter + " as " + anonym_datacenter[vDatacenter])

        vCluster = str(row['Cluster']).strip()
        # if Cluster not in anonymization dictionary add it although this could mean corrupted data
        if anonym_cluster.get(vCluster) is None:
            seqCluster = seqCluster + 1
            warns.append(vCluster + " cluster exists in vInfo but not in vCluster !! ")
            anonym_cluster[vCluster] = "Cluster-" + '{n:03}'.format(n=seqCluster) + "(!)"
            # print("Anonym: "+vDatacenter+"/"+vCluster+" as "+anonym_datacenter[vDatacenter]+"/"+anonym_cluster[vCluster])

        # Update df
        vInfo_df.at[index, "Datacenter"] = anonym_datacenter[vDatacenter]
        vInfo_df.at[index, "Cluster"] = anonym_cluster[vCluster]

    ########################################################
    if debug: print("Anonymize vHost: Datacenter Cluster")  #
    ########################################################
    vHost_df = combine_data_sheets(sheets_dict, 'vHost')

    # Iterate vInfo for Datacenter and Cluster
    for index, row in vHost_df.iterrows():
        vCluster = str(row['Cluster']).strip()
        vDatacenter = str(row['Datacenter']).strip()

        if anonym_cluster.get(vCluster) is None:
            warns.append(vCluster + " cluster exists in vHost but not in vCluster !! ")
        else:
            vHost_df.at[index, "Cluster"] = anonym_cluster[vCluster]

        if anonym_datacenter.get(vDatacenter) is None:
            warns.append('"' + vDatacenter + '" datacenter exists in vHost but not in vInfo !! ')
        else:
            vHost_df.at[index, "Datacenter"] = anonym_datacenter[vDatacenter]

    ########################################################
    if debug: print("Anonymize vNIC: Datacenter Cluster")  #
    ########################################################
    vNIC_df = combine_data_sheets(sheets_dict, 'vNIC')

    # Iterate vInfo for Datacenter and Cluster
    for index, row in vNIC_df.iterrows():
        vCluster = str(row['Cluster']).strip()
        vDatacenter = str(row['Datacenter']).strip()

        if anonym_cluster.get(vCluster) is None:
            warns.append(vCluster + " cluster exists in vNIC but not in vCluster !! ")
        else:
            vNIC_df.at[index, "Cluster"] = anonym_cluster[vCluster]

        if anonym_datacenter.get(vDatacenter) is None:
            warns.append('"' + vDatacenter + '" datacenter exists in vNIC but not in vInfo !! ')
        else:
            vNIC_df.at[index, "Datacenter"] = anonym_datacenter[vDatacenter]

    data["datacenter"] = anonym_datacenter
    data["cluster"] = anonym_cluster

    return data


'''
########################################################################################################
# PRINT FUNCTIONS
########################################################################################################
'''
def print_versions(sheets_dict):
    # see: https://www.robware.net/readMore

    add_p().add_run("AnalytiX build: " + str(version) + " - © Red Hat, 2025")

    try:
        metadata = get_rows(sheets_dict, sheet='vMetaData', key_columns='RVTools_major_version',
                            columns=['RVTools_major_version', 'xlsx_creation_datetime'], debug=False)
        if metadata.size > 0:
            p = add_p()
            p.add_run("RVTools Version:" + str(metadata['RVTools_major_version'].iloc[0]) + "  Creation Date:" + str(
                metadata['xlsx_creation_datetime'].iloc[0]))
        else:
            add_p().add_run(
                'There is no information about RVTools Version and XLS creation date: "Empty vMetaData sheet"').font.color.rgb = red_color
    except Exception as ex:
        add_p().add_run("Error retrieving RVTools version and XLS creation date from vMetaData: " + str(
            ex)).font.color.rgb = red_color

    # VERSIONS
    # #####################################################################################################################
    # ver: https://www.robware.net/readMore

    p = add_p()
    try:
        systemIDs = get_rows(sheets_dict, sheet='vSource', key_columns='Fullname', columns=["Fullname", "API_version"],
                             debug=False)
        supportedAPI = "6.5"
        for row in systemIDs.itertuples():
            p.add_run("API Connection: " + row.Fullname)
            if str(row.API_version) < supportedAPI:
                run = p.add_run(" (unsupported: Version " + supportedAPI + " or later is required)")
                run.font.color.rgb = red_color
            run = p.add_run("\n")
    except Exception as ex:
        p.add_run('ERROR: Can\'t check API Version: ' + str(ex) + '\n').font.color.rgb = red_color
        p.add_run('Column "API Version" does not exist in vSource' + '\n')
        p.add_run('Check the API Version is equal or greater than "6.5"' + '\n')

    return


def print_initialize(document):
    style = document.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10)

    sections = document.sections
    for section in sections:
        # Margins
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2)
        section.right_margin = Cm(1)
        # Header
        # section.header_distance = Mm(12.7)
        # section.footer_distance = Mm(12.7)
        # A4
        section.page_height = Mm(297)
        section.page_width = Mm(210)
    return


def print_vmw_products(sheets_dict):
    try:
        licenses = get_rows(sheets_dict, sheet='vLicense', key_columns='Name', columns=["Name"], debug=False)
        if licenses.size > 0:
            table_from_df(licenses, document, True, [500], "VMWare Installed Products")
            add_p().add_run("")
        else:
            add_p().add_run(
                'There is no VMWare installed products information: "Empty vLicense sheet"').font.color.rgb = red_color
            print("vLicense WARNING: vLicense tab not found in spreadsheet. Report will be partially generated")
            print("vLicense WARNING: Please add it to RVTools output configuration for better analysis results")

    except Exception as ex:
        add_p().add_run("Error retrieving VMWare installed products information from vLicense: " + str(
            ex)).font.color.rgb = red_color

    # Calculate ESX percentages and print the results
    # #####################################################################################################################

    ESX_data = calculate_percentage(sheets_dict, sheet='vHost', columns='ESX_Version')
    add_percentage_table(document, ESX_data, "ESX_Version", None, "Table Grid", None, None)
    add_p().add_run("")

    # VMWare Tools Status
    # #####################################################################################################################
    # Calculate VMWare Tools status percentages and print the results
    # https://docs.redhat.com/en/documentation/migration_toolkit_for_virtualization/2.5/html/installing_and_using_the_migration_toolkit_for_virtualization/prerequisites#vmware-prerequisites_mtv
    # TODO: check if this is required to validate some specific or minimum VMWare Tools version
    tools_data = calculate_percentage(sheets_dict, sheet='vTools', columns='Tools')
    rowColor = []
    for value, count, percentage in tools_data:
        if value == "toolsOk":
            rowColor.append(green_color)
        elif value == "toolsOld":
            rowColor.append(blue_color)
        elif value == "toolsNotRunning":
            rowColor.append(orange_color)
        elif value == "toolsNotInstalled":
            rowColor.append(red_color)
        else:
            rowColor.append(red_color)
            value = ""

    add_percentage_table(document, tools_data, "[vTools] Tools (VMWare-Tools status)", rowColor, "Table Grid")
    add_p().add_run("")
    return


def print_compute_sizing(sheets_dict):
    # Host Servers COMPUTE sizing
    # #####################################################################################################################

    vHostCores = groupby(sheets_dict, sheet='vHost', columns=["N_CPU", "Cores_per_CPU", "N_Cores", "CPU_Model", "Host"],
                         ascending=False, debug=False)
    table = table_from_df(vHostCores, document, grid=True, columnWidths=[50, 90, 60, 250, 50], title="Hardware")

    # calculate totals and Set RED color for processors with less than 8 cores.
    totCPU = 0
    avgCoresCPU = 0
    totCores = 0
    totSPairs = 0
    for row in table.rows:
        if row.cells[1].text.isnumeric():
            totCPU = totCPU + int(row.cells[0].text) * int(row.cells[4].text)
            totCores = totCores + int(row.cells[2].text) * int(row.cells[4].text)
            socks = int(row.cells[0].text)
            if (socks == 1): socks = 2  # manage hosts with just 1 socket
            totSPairs = totSPairs + (socks / 2) * int(row.cells[4].text)
            if int(row.cells[1].text) < 8:
                for cell in row.cells: cell.paragraphs[0].runs[0].font.color.rgb = red_color

    # print(str(totCPU))
    # print(str(totCores))
    add_p().add_run(
        "Total CPUs:" + str(totCPU) + "     Avg.Cores/CPU:" + str(int(totCores / totCPU)) + "     Total Cores:" + str(
            totCores) + "     Total Socket Pairs:" + str(int(totSPairs)))
    # addP().add_run("")

    # Host server models
    # #####################################################################################################################
    rowLink = []
    dataList = calculate_percentage(sheets_dict, sheet='vHost', columns=['Vendor', 'Model'])

    for value, count, percentage in dataList:
        stop = value[1].find('-')  # Facilitate finding LENOVO models usually having extra text after -
        if stop > 0:
            searchStr = str(value[1])[0:stop]
        else:
            searchStr = value[1]
        searchStr = searchStr.replace(':', '')  # Usually found an undesired ":" after model
        # print("value[1]=" + value[1] + "     slash=" + str(stop) + "    searchStr=" + searchStr)
        rowLink.append(
            "https://catalog.redhat.com/search?gs&q=" + urllib.parse.quote(searchStr, safe='/', encoding=None,
                                                                           errors=None) + "&target_platforms=Red Hat OpenShift")

    add_percentage_table(document, dataList, "[vHost] Vendor/Model", None, "Table Grid", None, rowLink,
                         title="Host servers models")
    add_p().add_run("")

    # CPU, VMs, Resource Pools... table
    # get data per datacenter/cluster from vHost (#VMs from vHost can be slightly different to count rows in vInfo)
    # #####################################################################################################################

    total_number_of_Datacenters = '0'
    total_number_of_Clusters = '0'
    total_number_of_Hosts = '0'
    total_number_of_CPUs = '0'
    total_number_of_Cores = '0'
    total_number_of_running_vCPU = '0'
    total_number_of_vCPU = '0'
    total_number_of_GiB = '0'
    total_number_of_VMs = '0'
    total_rpools_per_cluster = '0'

    hosts_per_cluster = groupby(sheets_dict, sheet='vHost', columns=["Datacenter", "Cluster", "N_Cores"], trunk=True,
                                result_name="Hosts", sum=False, debug=False)
    # print("===== hosts_per_cluster =====")
    # print(hosts_per_cluster)
    cpu_per_cluster = groupby(sheets_dict, sheet='vHost', columns=["Datacenter", "Cluster", "N_CPU"], trunk=True,
                              result_name="CPUs", sum=True, debug=False)
    # print("===== cpu_per_cluster =====")
    # print(cpu_per_cluster)
    cores_per_cluster = groupby(sheets_dict, sheet='vHost', columns=["Datacenter", "Cluster", "N_Cores"], trunk=True,
                                result_name="Total\nCores", sum=True, debug=False)
    # print("===== cores_per_cluster =====")
    # print(cores_per_cluster)
    running_vcpu_per_cluster = groupby(sheets_dict, sheet='vHost', columns=["Datacenter", "Cluster", "N_vCPUs"],
                                       trunk=True, result_name="Run\nvCPUs", sum=True, debug=False)
    # print("===== running_vcpu_per_cluster =====")
    # print(running_vcpu_per_cluster)
    vcpu_per_cluster = groupby(sheets_dict, sheet='vCPU', columns=["Datacenter", "Cluster", "CPUs"], trunk=True,
                               result_name="Total\nvCPUs", sum=True, debug=False)
    # print("===== vcpu_per_cluster =====")
    # print(vcpu_per_cluster)
    mem_per_cluster = groupby(sheets_dict, sheet='vInfo', columns=["Datacenter", "Cluster", "Memory"], trunk=True,
                              result_name="Mem\nGiB", sum=True, debug=False)

    # print("===== maxmem_per_cluster =====")
    # print(maxmem_per_cluster)

    vms_per_cluster = groupby(sheets_dict, sheet='vHost', columns=["Datacenter", "Cluster", "N_VMs_total"], trunk=True,
                              result_name="VMs", sum=True, debug=False)

    rpools_per_cluster = get_rows(sheets_dict, sheet='vInfo', key_columns=["Datacenter", "Cluster", "Resource_pool"],
                                  columns=["Datacenter", "Cluster", "Resource_pool"], debug=False)
    rpools_per_cluster = groupby_df(rpools_per_cluster, columns=["Datacenter", "Cluster", "Resource_pool"], trunk=True,
                                    result_name="Res.\nPools", debug=False)
    # tot_rps = rpools_per_cluster["RPs"].sum(numeric_only=True)
    total_rpools_per_cluster = 0
    for index, row in rpools_per_cluster.iterrows(): total_rpools_per_cluster = total_rpools_per_cluster + row[
        "Res.\nPools"]

    # Merge
    merged_df = pd.merge(hosts_per_cluster, cpu_per_cluster, on=["Datacenter", 'Cluster'])
    merged_df = pd.merge(merged_df, cores_per_cluster, on=["Datacenter", 'Cluster'])
    merged_df = pd.merge(merged_df, running_vcpu_per_cluster, on=["Datacenter", 'Cluster'])
    merged_df = pd.merge(merged_df, vcpu_per_cluster, on=["Datacenter", 'Cluster'])
    merged_df = pd.merge(merged_df, mem_per_cluster, on=["Datacenter", 'Cluster'])
    merged_df = pd.merge(merged_df, vms_per_cluster, on=["Datacenter", 'Cluster'])
    merged_df = pd.merge(merged_df, rpools_per_cluster, on=["Datacenter", 'Cluster'])
    # print("===== merged_df =====")
    # print(merged_df)

    oversubs_ratio = []  # oversubscription ratio for each cluster and Memory MiB > Gib
    for index, row in merged_df.iterrows():
        # CPU Oversubscription
        oversubs_ratio.append(round(row["Total\nvCPUs"] / row["Total\nCores"] / 2, 2))
        # print("oversubs_ratio["+str(index)+"] = "+ str(row["Total\nvCPUs"]) + " / " + str(row["Total\nCores"]) + " / 2 = " + str(oversubs_ratio[index]))

        # MEMORY  MiB >> GiB
        merged_df.at[index, 'Mem\nGiB'] = int(int(row['Mem\nGiB']) / 1000)

    merged_df.insert(7, "CPU\nOvers", oversubs_ratio, True)
    table = table_from_df(merged_df, document, grid=True, columnWidths=[80, 120, 34, 32, 34, 36, 36, 35, 35, 30, 32],
                          title="[vHost/vCPU] Hosts per Cluster")

    # Alternate background color for datacenters and reduce font
    lastDC = ""
    count = 0
    for row in table.rows:
        if count > 1:  # Reduce data rows font to 8 by default

            for cell in row.cells:
                cell.paragraphs[0].runs[0].font.size = Pt(8)
                set_cell_margins(cell, start=0, end=20)

            if len(row.cells[0].text) > 20:
                row.cells[0].paragraphs[0].runs[0].font.size = Pt(6)
            else:
                row.cells[0].paragraphs[0].runs[0].font.size = Pt(7)
            if len(row.cells[1].text) > 32:
                row.cells[1].paragraphs[0].runs[0].font.size = Pt(6)
            else:
                row.cells[1].paragraphs[0].runs[0].font.size = Pt(7)

        if count == 1:  # Reduce headers font
            for cell in row.cells:
                cell.paragraphs[0].runs[0].font.size = Pt(9)

        if row.cells[0].text != lastDC:  # Alternate background
            count += 1
            lastDC = row.cells[0].text

        if count % 2 == 0:
            for cell in row.cells: set_cell_background(cell, "E5E5EE")

    row = table.add_row()
    total_number_of_Datacenters = str(count_rows(sheets_dict, 'vInfo', key_columns='Datacenter', debug=False))
    total_number_of_Clusters = str(count_rows(sheets_dict, 'vCluster'))
    total_number_of_Hosts = str(count_rows(sheets_dict, 'vHost'))
    total_number_of_CPUs = str(sum_rows(sheets_dict, 'vHost', 'N_CPU', debug=False))
    total_number_of_Cores = str(sum_rows(sheets_dict, 'vHost', 'N_Cores', debug=False))
    total_number_of_running_vCPU = str(sum_rows(sheets_dict, 'vHost', 'N_vCPUs', debug=False))
    total_number_of_vCPU = str(
        int(sum_rows(sheets_dict, 'vCPU', 'CPUs', debug=False)))  # int( because sometimes found floats :-0
    total_number_of_GiB = str(int(sum_rows(sheets_dict, 'vMemory', 'Max', debug=False) / 1000))
    total_number_of_VMs = str(sum_rows(sheets_dict, 'vHost', 'N_VMs_total', debug=False))

    row.cells[0].text = total_number_of_Datacenters
    row.cells[1].text = total_number_of_Clusters
    row.cells[2].text = total_number_of_Hosts
    row.cells[3].text = total_number_of_CPUs
    row.cells[4].text = total_number_of_Cores
    row.cells[5].text = total_number_of_running_vCPU
    row.cells[6].text = total_number_of_vCPU
    row.cells[7].text = str(round(int(total_number_of_vCPU) / int(total_number_of_Cores) / 2, 2))
    row.cells[8].text = total_number_of_GiB
    row.cells[9].text = total_number_of_VMs
    row.cells[10].text = str(total_rpools_per_cluster)

    for cell in row.cells:
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        cell.paragraphs[0].runs[0].font.bold = True
        if len(cell.text) > 6:
            cell.paragraphs[0].runs[0].font.size = Pt(6)
        elif len(cell.text) == 6:
            cell.paragraphs[0].runs[0].font.size = Pt(7)
        else:
            cell.paragraphs[0].runs[0].font.size = Pt(9)
        set_cell_margins(cell, start=0, end=20)

    add_p().add_run("Oversubscription Ratio = (Total vCPUs / Total Cores / 2)")

    return row


def print_memory_ranks(sheets_dict):
    ################################################################################
    # VM's Memory List
    ################################################################################

    # Retrieve information
    memory_config = get_rows(sheets_dict, 'vInfo', key_columns=None, columns=['VM', 'Memory'])
    memory_config = memory_config.sort_values(by=['Memory'], ascending=False)

    lmemory = []
    pre_vm = ''
    pre_memory = ''
    vm = ''
    count = 0

    # Add headers
    lmemory.append(["VMs", "GiB", "Cnt"])

    # Add content
    for index, row in memory_config.iterrows():

        memory = f'{int(int(row["Memory"]) / 1024):n}'
        vm = row['VM']

        if pre_memory != memory:
            if pre_vm != '':  # do not add the first empty element
                lmemory.append([pre_vm, pre_memory, count])
            pre_vm = ''
            count = 1
        else:
            count += 1

        pre_memory = memory
        if pre_vm == '':
            pre_vm = str(vm)
        else:
            pre_vm = pre_vm + ", " + str(vm)

    lmemory.append([pre_vm, pre_memory, count])

    # Write to DOC
    if memory_config is not None:
        add_p().add_run("")

        L = WD_ALIGN_PARAGRAPH.LEFT
        R = WD_ALIGN_PARAGRAPH.RIGHT
        table = table_from_list(lmemory, document, True, [440, 30, 30], [L, R, R],
                                "[vInfo] VMs grouped by memory size")

        # Adjust fonts and lengths
        # 0:VMs - 1:MemoryCapacity - 2:Annotations - 3:Count
        count = 0
        for row in table.rows:
            count += 1
            if count > 2:  # jump headers
                for cell in row.cells: cell.paragraphs[0].runs[0].font.size = Pt(8)
                max_len = 250
                if len(str(row.cells[0].text)) > max_len: row.cells[0].text = row.cells[0].text[0:max_len] + "..."
                row.cells[0].paragraphs[0].runs[0].font.size = Pt(6)

    ################################################################################
    # Hosts Memory List
    ################################################################################

    # Retrieve information
    memory_config = get_rows(sheets_dict, 'vHost', key_columns=None, columns=['Host', 'N_Memory'])
    memory_config = memory_config.sort_values(by=['N_Memory'], ascending=False)
    # print("====================================")
    # print(memory_config)

    lmemory = []
    pre_host = ''
    pre_memory = ''
    host = ''
    count = 0

    # Add headers
    lmemory.append(["Hosts", "GiB", "Cnt"])

    # Add content
    for index, row in memory_config.iterrows():
        memory = f'{int(int(row["N_Memory"]) / 1024):n}'
        host = row['Host']

        # print("memory:"+str(memory)+"    host:"+host)
        if pre_memory != memory:
            if pre_memory != '':  # do not add the first empty element
                lmemory.append([pre_host, pre_memory, count])
            pre_host = ''
            count = 1
        else:
            count += 1

        pre_memory = memory
        if pre_host == '':
            pre_host = str(host)
        else:
            pre_host = pre_host + ", " + str(host)

    lmemory.append([pre_host, pre_memory, count])

    # Write to DOC
    if memory_config is not None:
        add_p().add_run("")

        L = WD_ALIGN_PARAGRAPH.LEFT
        R = WD_ALIGN_PARAGRAPH.RIGHT
        table = table_from_list(lmemory, document, True, [430, 40, 30], [L, R, R],
                                "[vHost] Hosts grouped by physical memory size")

        # Adjust fonts and lengths
        # 0:VMs - 1:MemoryCapacity - 2:Annotations - 3:Count
        count = 0
        for row in table.rows:
            count += 1
            if count > 2:  # jump headers
                for cell in row.cells: cell.paragraphs[0].runs[0].font.size = Pt(8)

                if len(str(row.cells[0].text)) > 300: row.cells[0].text = row.cells[0].text[0:300] + "..."
                row.cells[0].paragraphs[0].runs[0].font.size = Pt(6)

    return


def print_compute_checks_and_hints(sheets_dict, sizing_row):
    # #####################################################################################################################
    # Compute Checks & Hints... table
    # #####################################################################################################################
    add_p().add_run("")

    compute_checks = {}
    key = "Clusters with less than 3 hosts (vCluster)"
    try:
        compute_checks[key] = count_rows(sheets_dict, 'vCluster', query_expr='NumHosts < 3')
    except Exception as ex:
        compute_checks[key] = str(ex)

    key = "VMs with CBT not activated (vInfo)"
    try:
        compute_checks[key] = count_rows(sheets_dict, "vInfo", query_expr="CBT!=True")
    except Exception as ex:
        compute_checks[key] = str(ex)

    key = "Hosts with more than 64 cores (vHost)"
    try:
        compute_checks[key] = count_rows(sheets_dict, 'vHost', query_expr='N_Cores > 64')
    except Exception as ex:
        compute_checks[key] = str(ex)

    key = "Hosts with more than 128 cores (vHost)"
    try:
        compute_checks[key] = count_rows(sheets_dict, 'vHost', query_expr='N_Cores > 128')
    except Exception as ex:
        compute_checks[key] = str(ex)

    key = "VMs with CPU HotAdd (vCPU)"
    try:
        compute_checks[key] = count_rows(sheets_dict, 'vCPU', query_expr='Hot_Add == True')
    except Exception as ex:
        compute_checks[key] = str(ex)

    # key = "VMs with CPU HotRemove (vCPU)"
    # try: compute_checks[key] = count_rows(sheets_dict, 'vCPU', query_expr='Hot_Remove == True')
    # except Exception as ex: compute_checks[key] = str(ex)

    # key = "VMs with CPU Numa HotAdd Exposed (vCPU)"
    # try: compute_checks[key] = count_rows(sheets_dict, 'vCPU', query_expr='Numa_Hotadd_Exposed == True')
    # except Exception as ex: compute_checks[key] = str(ex)

    # key = "VMs with RAM btw 1TB and 2TB (vInfo)"
    # try: compute_checks[key] = count_rows(sheets_dict, 'vInfo', query_expr='Memory > 100000 and Memory < 200000 ')
    # except Exception as ex: compute_checks[key] = str(ex)

    key = "VMs with RAM btw 2TB and 6TB (vInfo)"
    try:
        compute_checks[key] = count_rows(sheets_dict, 'vInfo', query_expr='Memory >= 200000 and Memory <= 600000 ')
    except Exception as ex:
        compute_checks[key] = str(ex)

    key = "VMs with RAM greater than 6TB (vInfo)"
    try:
        compute_checks[key] = count_rows(sheets_dict, 'vInfo', query_expr='Memory > 600000')
    except Exception as ex:
        compute_checks[key] = str(ex)

    key = "VMs with memory HotAdd (vMemory)"
    try:
        compute_checks[key] = count_rows(sheets_dict, 'vMemory', query_expr='Hot_Add == True')
    except Exception as ex:
        compute_checks[key] = str(ex)

    key = "VMs with Ballooned memory (vMemory)"
    try:
        compute_checks[key] = count_rows(sheets_dict, 'vMemory', query_expr='Ballooned == True')
    except Exception as ex:
        compute_checks[key] = str(ex)

    # vCluster Num VMotions
    # vCluster DRS vmotion rate
    # vHost VMotion support
    # vHost Storage VMotion support

    table = table_from_dict(compute_checks, document, True, None,
                            "Check", "Data", True, "Compute Checks & Hints")

    # TODO: Avoid this ugly and crazy way to reuse information
    total_number_of_Clusters = sizing_row.cells[1].text
    total_number_of_Hosts = sizing_row.cells[2].text
    total_number_of_VMs = sizing_row.cells[9].text

    run = add_p().add_run("Reference:  Total Clusters=" + total_number_of_Clusters +
                         "   Total Hosts=" + total_number_of_Hosts +
                         "   Total VMs=" + total_number_of_VMs)
    run.font.color.rgb = gray_color

    return


def print_networking(sheets_dict):
    #######################################################################################################################
    # %%% REPORT Networking

    add_p().add_run("")

    # Show the NIC drivers by datacenter and link status
    # #####################################################################################################################
    '''
    network_drivers = groupby(sheets_dict, sheet='vNIC', columns=["Datacenter", "Driver", "Speed", "Duplex", "MAC"], debug=False)
    table = table_from_df(network_drivers, document, grid=True, columnWidths=[75,75,50,75,50],
                          title="[vNIC] NIC drivers by datacenter and link status")
    i=0
    for row in table.rows:
        #TODO: there should be a way to do it in table_from_df as we do with percentage
        if i > 1: #avoid table headers
            #add link
            paragraph=row.cells[1].paragraphs[0]
            text = row.cells[1].paragraphs[0].text
            row.cells[1].paragraphs[0].text = ""
            if text == "cdce":
                link = "https://www.google.com/search?q=cdce+vusb0" # cdce is USB interface
                text = "cdce (usb?)"
            else: link = "https://www.vmware.com/resources/compatibility/vcl/result.php?search="+text+"&searchCategory=io"
            run = add_hyperlink(paragraph, text, link)
            run.font.underline = True
            run.font.color.rgb = blue_color
            #speed in Gb
            text = row.cells[2].paragraphs[0].text
            speed=float(text)
            row.cells[2].paragraphs[0].text = ""
            speed=speed/1000
            if (speed < 1) : row.cells[2].paragraphs[0].text = "{:.1f} Gb".format(speed)
            else: row.cells[2].paragraphs[0].text = "{:.0f} Gb".format(speed)
            row.cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            row.cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        i += 1

    for row in table.rows:
        textlist = row.cells[0].text.split("-")
        if len(textlist) > 1 and textlist[0] == "Datacenter" and int(textlist[1]) % 2 == 0 :
            for cell in row.cells: set_cell_background(cell,"E5E5EE")

    addP().add_run("")
    '''
    # Show the number of NICs per host
    # #####################################################################################################################

    nics_host_shortList = {}
    NIC_Host = groupby(sheets_dict, sheet='vNIC', columns=["Host", "MAC"], ascending=False, debug=False)
    NIC_Host = NIC_Host.groupby(["Count"]).count().reset_index().sort_values(by=['Count'], ascending=True)

    set_run_title_style(add_p().add_run("[vNIC] NIC/Host (NICs per Host)"))
    p = add_p()
    for index, row in NIC_Host.iterrows():
        p.add_run("There are " + str(row["Host"]) + " hosts with " + str(row["Count"]) + " MAC address (NICs)\n")

    # Show the number of NICs per VM
    # #####################################################################################################################
    nics_vm_shortList = {}
    NIC_VM = groupby(sheets_dict, sheet='vInfo', columns=["NICs", "VM"]).sort_values(by=['NICs'], ascending=True)
    # print("NICs per VM")
    # print(NIC_VM)

    set_run_title_style(add_p().add_run("[vInfo] NIC/VM (NICs per VM)"))
    p = add_p()
    for index, row in NIC_VM.iterrows():
        p.add_run("There are " + str(row["Count"]) + " VMs with " + str(row["NICs"]) + " NICs\n")

    # Show the number of VMs attached to each network
    # #####################################################################################################################

    network_VM_shortList = {}
    set_run_title_style(add_p().add_run("[vNetwork] Network/VM (VMs per Network)"))
    network_VM = groupby(sheets_dict, sheet='vNetwork', columns=["Network", "VM"], ascending=False, debug=False)
    network_VM = network_VM.groupby(["Count"]).count().reset_index().sort_values(by=['Count'], ascending=True)

    for row in network_VM.itertuples():
        if (row.Count < 5):
            index = "just " + str(row.Count)
            currentValue = network_VM_shortList.get(index, 0)
            newValue = currentValue + row.Network
            network_VM_shortList.update({index: newValue})
        if row.Count >= 5 and row.Count < 10:
            index = ">=5..<10"
            currentValue = network_VM_shortList.get(index, 0)
            newValue = currentValue + row.Network
            network_VM_shortList.update({index: newValue})
        if row.Count >= 10 and row.Count < 50:
            index = ">=10..<50"
            currentValue = network_VM_shortList.get(index, 0)
            newValue = currentValue + row.Network
            network_VM_shortList.update({index: newValue})
        if row.Count >= 50 and row.Count < 100:
            index = ">=50..<100"
            currentValue = network_VM_shortList.get(index, 0)
            newValue = currentValue + row.Network
            network_VM_shortList.update({index: newValue})
        if row.Count > 100:
            index = ">=100"
            currentValue = network_VM_shortList.get(index, 0)
            newValue = currentValue + row.Network
            network_VM_shortList.update({index: newValue})

    p = add_p()
    for key, value in network_VM_shortList.items():
        p.add_run("There are " + str(value) + " networks with " + key + " VMs attached\n")

    return


def print_network_terms(sheets_dict):
    # Search Network related terms
    # #####################################################################################################################

    terms = {"Cisco": "cisco",
             "EPG": "\\bepg\\b|_epg|epg_",
             "ACI": "\\baci\\b|_aci|aci_",
             "NSX": "nsx",
             "Infoblox": "nios|infoblox",
             "F5": "big-iq|f5.com",
             "NetScout": "ngeniusone|netscout"
             }

    sheets = ['vNetwork', 'vNIC', 'vSwitch', 'dvSwitch', 'vPort', 'dvPort']

    headers = [["Search Term", "Count", "Column", "Spreadsheet"]]

    results = global_search(sheets_dict, terms, sheets)
    network_findings = headers + results

    tb = table_from_list(network_findings, document, columns_width=[300, 50, 75, 75], grid=True,
                         title="Network related terms in " + str(sheets))
    # title="Search of network related terms in vNetwork,vNIC,vSwitch,vPort and dvSwitch sheets")
    add_p().add_run("")
    return


def print_workload_terms(sheets_dict):
    # Search workload related terms
    # #####################################################################################################################
    # search examples:
    # contains a string: 'string'
    # contains string1 or string2: 'string1|string2'
    # contains string1 and string2: '(?=.*string1)(?=.*string2)'
    # contains a word: '\\bword\\b'
    # contains word1 or word2: '\\bword1\\b|\\bword2\\b'
    # contains word1 and word2: '(?=.*\\bstring1\\b)(?=.*\\bstring2\\b)'
    # contains "mariadb" "maria-db" or "maria db": 'maria.?db'

    terms = {"Ansible": "ansible",
             "JBoss": "jboss",
             "tomcat": "tomcat",
             "Weblogic": "weblogic",
             "Websphere": "websphere",
             "Spring": "spring",
             "Apache": "apache",
             "HA-Proxy": "ha.?proxy",
             "Apigee": "apigee",
             "Jenkins": "jenkins",
             "OpenShift": "openshift",
             "Oracle": "oracle",
             "Postgre": "postgre",
             "Kafka": "kafka",
             "Splunk": "splunk",
             "Spark": "spark",
             "Mysql": "mysql",
             "Informix": "informix",
             "MariaDB": "maria.?db",
             "MongoDB": "mongo",
             "Redis": "redis",
             "MS-Sql": "sql.?server|always on",
             "Lotus": "lotus",
             "SAP": "\\bsap\\b|hana",
             "Documentum": "documentum",
             "Vignette": "vignette",
             "Filenet": "filenet",
             "SAS": "\\bsas\\b|/sas/",
             "palo.?alto": "palo.?alto",
             "qualys": "qualys",
             "cyberark": "cyberark",
             "forti.?net|fortiauth": "forti.?net|fortiauth",
             "checkpoint": "checkpoint",
             "open.?am|open.?idm": "open.?am|open.?idm",
             "pandora": "pandora",
             "datadog": "datadog",
             "netskope": "netskope",
             "Dynatrace": "dynatrace",
             "one.?view": "one.?view",
             "Azure": "\\barc\\b|azure",
             "AWS": "\\baws\\b|amazon",
             "NetApp": "netapp|ontap",
             "PureStorage": "pure.?storage|portworx",
             "Data Protector": "data.?protector",
             "Dell PowerProtect": "powerprotect",
             "Avamar": "avamar",
             "Veritas": "veritas|netbackup",
             "Veeam": "veeam",
             "Commvault": "commvault",
             "IBM Storage Protect": "storage.?protect",
             "Storware": "Storware",
             "Trilio": "trilio",
             "Citrix": "citrix",
             "VDI": "\\bvdi\\b|\\bhorizon\\b",
             }

    # the sheets we are going to search in
    sheets = ['vNetwork', 'vNIC', 'vSwitch', 'vPort', 'dvSwitch', 'vPartition']

    headers = [["Search Term", "Count", "Column", "Spreadsheet"]]

    results = global_search(sheets_dict, terms, sheets)

    results = [tup for tup in results if tup[1] > 5 and "OS_according_to_" not in tup[2]]
    network_findings = headers + results

    tb2 = table_from_list(network_findings, document, columns_width=[300, 50, 75, 75], grid=True,
                          title="Workload related terms in " + str(sheets))

    add_p().add_run("")
    return


def print_appliances_n_OVA_annotations(sheets_dict):
    # Search for appliance or OVA in annotations contains_expr='appliance|\\bova\\b'
    # #####################################################################################################################
    appliances = get_rows(sheets_dict, 'vInfo', debug=False, contains_column='Annotation',
                          contains_expr='appliance|\\bova\\b', contains_case=False)
    groupedAppliances = groupby_df(appliances, columns=["Annotation", "VM"], debug=False)
    groupedAppliances = groupedAppliances.sort_values(by=['Count'], ascending=False)
    tb = table_from_df(groupedAppliances, document, grid=True, columnWidths=[450, 50],
                       title="[vInfo] Annotation filtered by 'Appliance and OVA'")

    for row in tb.rows:
        if len(str(row.cells[0].text)) > 180:
            row.cells[0].text = row.cells[0].text[0:180] + "..."

    # tb.cell(0, 1).text = 'Count'
    # tb.cell(0, 1).paragraphs[0].runs[0].font.bold = True

    add_p().add_run("")

    return


def print_operating_systems(sheets_dict):
    # List all Operating Systems with its supportability status
    # https://access.redhat.com/articles/4234591
    # #####################################################################################################################

    supported_OS = [
        'Red Hat Enterprise Linux 9 (64-bit)'
        , 'Red Hat Enterprise Linux 8 (64-bit)'
        , 'Red Hat Enterprise Linux 7 (64-bit)'
        , 'Microsoft Windows Server 2022 (64-bit)'
        , 'Microsoft Windows Server 2019 (64-bit)'
        , 'Microsoft Windows Server 2016 or later (64-bit)'
        , 'Microsoft Windows Server 2016 (64-bit)'
        , 'Microsoft Windows 10 (64-bit)'
        , 'SUSE Linux Enterprise 15 (64-bit)'
    ]

    EOL_OS = [
        'Red Hat Enterprise Linux 6 (64-bit)'
        , 'Microsoft Windows Server 2012 (64-bit)'
        , 'Ubuntu Linux (64-bit)'
    ]

    Community_OS = [
        'CentOS 7 (64-bit)'
        , 'CentOS 8 (64-bit)'
        , 'CentOS 9 (64-bit)'
        , 'Red Hat Fedora (64-bit)'
    ]

    # Repair empty cells. get_rows without columns to allow update
    os_data_df = get_rows(sheets_dict, sheet='vInfo')
    for index, row in os_data_df.iterrows():
        os = row['OS_according_to_the_VMware_Tools']
        if os is None or str(os).upper() == 'NAN' or (isinstance(os, str) and os == ''):
            # Update
            os_data_df.at[index, 'OS_according_to_the_VMware_Tools'] = row['OS_according_to_the_configuration_file']

    # Calculate OS percentages and print the results
    os_data = calculate_percentage(sheets_dict, sheet='vInfo', columns='OS_according_to_the_VMware_Tools')

    supported = float(0);
    eol = float(0);
    comm = float(0);
    never = float(0)
    rowColor = []
    for value, count, percentage in os_data:
        if value in supported_OS:
            rowColor.append(green_color); supported += percentage
        elif value in EOL_OS:
            rowColor.append(blue_color); eol += percentage
        elif value in Community_OS:
            rowColor.append(gray_color); comm += percentage
        else:
            rowColor.append(red_color); never += percentage

    table = add_percentage_table(document, os_data, "Operating System", rowColor, "Table Grid",
                                 title="[vInfo] OS_according_to_the_VMware_Tools\n")
    p = table.rows[0].cells[0].paragraphs[0]
    p.add_run("Supported {:5.2f}".format(supported) + "%").font.color.rgb = green_color;
    p.add_run(" - ")
    p.add_run("End Of Life {:5.2f}".format(eol) + "%").font.color.rgb = blue_color;
    p.add_run(" - ")
    p.add_run("Community support {:5.2f}".format(comm) + "%").font.color.rgb = gray_color;
    p.add_run(" - ")
    p.add_run("Not certified {:5.2f}".format(never) + "%").font.color.rgb = red_color

    add_p().add_run("\n")
    return

    # ---Image example---
    '''
    # generate picture example
    memfile = BytesIO()
    plt.plot(network_VM2["Count"], network_VM2["Network"])
    plt.xlabel('Number of VMs per Network')
    plt.ylabel('Frequency')
    plt.savefig(memfile)

    # add picture to table
    cell=table.rows[0].cells[1]
    cell.width = w1
    cell.vertical_alignment = WD_ALIGN_VERTICAL.BOTTOM
    set_cell_margins(cell, start=0, end=0, top=0, bottom=0)

    pi=cell.add_paragraph()
    pi.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pi.vertical_alignment = WD_ALIGN_VERTICAL.BOTTOM
    pi.add_run().add_picture(memfile,width=w1)

    memfile.close() #close image file 
    '''


def print_storage_models(sheets_dict):
    dataList = calculate_percentage(sheets_dict, sheet='vMultiPath', columns=['Display_name', 'Vendor', 'Model'])
    rowLink = []
    for value, count, percentage in dataList:
        rowLink.append(
            "https://www.google.com/search?q=" + urllib.parse.quote(str(value[1]) + str(value[2]) + " CSI", safe='/',
                                                                    encoding=None, errors=None))

    add_percentage_table(document, dataList, "[vMultipath] DisplayName/Vendor/Model", None, "Table Grid", links=rowLink,
                         title="Number of Datastore connections (from Host) per storage system type. 1:n")
    add_p().add_run("")
    return


def print_storage_capacity(sheets_dict):
    # Datastores per type including capacity
    # #####################################################################################################################
    try:
        datastores = get_rows(sheets_dict, sheet='vDatastore',
                              columns=["Type", "Display_name", "Capacity_MiB", "Provisioned_MiB", "In_Use_MiB",
                                       "Object_ID"], debug=False)
        datastores = datastores.sort_values(by=["Type", "Display_name"], key=lambda x: x.str.lower())

        Type = ""
        Display = None

        Capacity = 0;
        Provisioned = 0;
        InUse = 0;
        Count = 0
        tCapacity = 0;
        tProvisioned = 0;
        tInUse = 0;
        tCount = 0

        L = WD_ALIGN_PARAGRAPH.LEFT
        R = WD_ALIGN_PARAGRAPH.RIGHT

        main_list = []
        t_main_list = []

        main_list.append(["Type", "Source", "Capacity", "Provsnd", "InUse", "Count"])

        columns_width = [45, 210, 65, 65, 65, 50]
        columns_align = [L, L, R, R, R, R]
        totals_text = "Type Totals in GiB"
        i = 0
        for index, row in datastores.iterrows():

            # print(row.Display_name)
            if Type != "" and Type.upper() != row.Type.upper():
                # when Type changes we save previous data on list
                t_main_list.append(
                    [Type, totals_text, f"{tCapacity / 1000:,.0f}", f"{tProvisioned / 1000:,.0f}",
                     f"{tInUse / 1000:,.0f}",
                     tCount])

                # initialize values at Type level
                tCapacity = 0;
                tProvisioned = 0;
                tInUse = 0;
                tCount = 0

            tCapacity = tCapacity + row.Capacity_MiB
            tProvisioned = tProvisioned + row.Provisioned_MiB
            tInUse = tInUse + row.In_Use_MiB
            tCount = tCount + 1

            if Display != None and str(Display).upper() != str(row.Display_name).upper():
                # print(Display + " != " + row.Display_name)
                # when Display changes we save previous data on list
                main_list.append(
                    [Type, Display, f"{Capacity / 1000:,.0f}", f"{Provisioned / 1000:,.0f}", f"{InUse / 1000:,.0f}",
                     Count])
                i += 1

                # initialize values at DisplayName level
                Capacity = 0;
                Provisioned = 0;
                InUse = 0;
                Count = 0

            Type = row.Type
            Display = row.Display_name
            Capacity += row.Capacity_MiB
            Provisioned += row.Provisioned_MiB
            InUse += row.In_Use_MiB
            Count += 1

        # Save last row values
        main_list.append(
            [Type, Display, f"{Capacity / 1000:,.0f}", f"{Provisioned / 1000:,.0f}", f"{InUse / 1000:,.0f}", Count])
        t_main_list.append(
            [Type, totals_text, f"{tCapacity / 1000:,.0f}", f"{tProvisioned / 1000:,.0f}", f"{tInUse / 1000:,.0f}",
             tCount])
        added_rows = len(t_main_list)
        tb = table_from_list(main_list + t_main_list, document, True, columns_width, columns_align,
                             title="[vDatastore] Datastores type, source (from address) and capacity in GiB")
        for x in range(added_rows):
            for y in range(len(tb.rows[0].cells)):
                cell = tb.cell(len(tb.rows) - x - 1, y)
                cell.paragraphs[0].runs[0].font.bold = True
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                set_cell_background(cell, "E5E5EE")
        p = add_p()
    except Exception as ex:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        p = add_p()
        p.add_run("[vDatastore] Datastores type, source (from address) and capacity in GiB")
        p.add_run("\nERROR processing section:" + str(ex))
        warns.append(
            "ERROR in section: [vDatastore] Datastores type, source (from address) and capacity in GiB (line:" + str(
                exc_tb.tb_lineno) + ") " + str(ex))

    return


def print_storage_connections(sheets_dict):
    """
    Print NFS, VMFS, VSAN storage in use by each Datacenter/Cluster
    Loop each DC/CL pair host in vHost. For each: check if storage exists in vDatastore and what type it is.

    :param sheets_dict:
    :return: generated table
    """
    # ####################################################################################################################
    hosts_per_cluster = get_rows(sheets_dict, sheet='vHost', key_columns=["Datacenter", "Cluster", "Host"],
                                 columns=["Datacenter", "Cluster", "Host"]).sort_values(by=["Datacenter", "Cluster"],
                                                                                        ascending=True)
    datastores = combine_data_sheets(sheets_dict, "vDatastore")

    matrix = []
    lastdc = "";
    lastcl = ""
    hasNFS = 0;
    hasVMFS = 0;
    hasVSAN = 0
    totNFS = 0;
    totVMFS = 0;
    totVSAN = 0

    #For each host
    for index, row in hosts_per_cluster.iterrows():
        dc = row["Datacenter"]
        cl = row["Cluster"]
        host = row["Host"]

        #subtotals per DC and Cluster
        if lastdc != dc or lastcl != cl:
            if lastdc != "": matrix.append([lastdc, lastcl, hasNFS, hasVMFS, hasVSAN])
            totNFS += hasNFS;
            totVMFS += hasVMFS;
            totVSAN += hasVSAN
            lastdc = dc;
            lastcl = cl;
            hasNFS = 0;
            hasVMFS = 0;
            hasVSAN = 0

        #datastores in use per current host
        found = datastores[datastores["Hosts"].str.contains(host, na=False, case=True)]
        for index2, dstore in found.iterrows():
            dstype = str(dstore["Type"]).upper()
            if dstype == "NFS": hasNFS += 1
            if dstype == "VMFS": hasVMFS += 1
            if dstype == "VSAN": hasVSAN += 1

    matrix.append(["", "", totNFS, totVMFS, totVSAN])
    matrixdf = pd.DataFrame(matrix, columns=["Datacenter", "Cluster", "NFS", "VMFS", "VSAN"])

    #generate document table
    tb = table_from_df(matrixdf, document, True, [190, 190, 40, 40, 40],
                       "[vHost/vDatastore] Host Datastore connections grouped by Cluster and Type")

    #adjust styles
    tb.cell(len(tb.rows) - 1, 2).paragraphs[0].runs[0].font.bold = True
    tb.cell(len(tb.rows) - 1, 3).paragraphs[0].runs[0].font.bold = True
    tb.cell(len(tb.rows) - 1, 4).paragraphs[0].runs[0].font.bold = True

    return tb


def print_storage_vms_per_disk_controller(sheets_dict):
    # [vDisk] Number of VMs using each type of Disk Controller"
    # ####################################################################################################################
    dataList = calculate_percentage(sheets_dict, sheet='vDisk', columns='Controller')
    add_p().add_run("")
    add_percentage_table(document, dataList, "vDisk Controller", None, "Table Grid",
                         title="[vDisk] Number of VMs using each type of Disk Controller")
    return


def print_storage_multiple_controllers(sheets_dict):
    # [vDisk] VMs with more than 1 disk controller
    # ####################################################################################################################
    add_p().add_run("")

    # Retrieve unique combinations of VM+Controller
    multiController_VM = get_rows(sheets_dict, sheet='vDisk', key_columns=["VM", "Controller"],
                                  columns=["VM", "Controller"])
    # Group to obtain the number of instances for each VM
    multiController_VM = groupby_df(multiController_VM, columns=["VM", "Controller"], trunk=True)
    # Filter VMs with more than one Controller
    controllersList = multiController_VM.query("Count > 1")
    # Print the list
    count = len(controllersList)
    if count > 0:
        query = ""
        first = True

        # build a query to retrieve all VM/Controller rows
        for index, row in controllersList.iterrows():
            if not first: query = query + " or "
            first = False
            query = query + "VM == '" + row["VM"] + "'"
        multiController_VM = get_rows(sheets_dict, sheet='vDisk', key_columns=["VM", "Controller"],
                                      columns=["VM", "Controller"], query_expr=query)
        count = len(multiController_VM)

        newValues = []
        for index, row in controllersList.iterrows():
            multiController_VM = get_rows(sheets_dict, sheet='vDisk', columns=["VM", "Controller"],
                                          query_expr="VM == '" + row["VM"] + "'")
            # print row.Controller
            ctrList = {}
            for index2, row2 in multiController_VM.iterrows():
                if row2.Controller in ctrList.keys():
                    ctrList[row2.Controller] = ctrList[row2.Controller] + 1
                else:
                    ctrList[row2.Controller] = 1

            newValues.append(str(len(multiController_VM)) + ": " + str(ctrList))

        controllersList = controllersList.drop('Count', axis=1)
        controllersList.insert(1, "Details", newValues, True)

        # Print the short list
        tb = table_from_df(controllersList, document, grid=True, columnWidths=[150, 350],
                           title="[vDisk] VM/Controller (VMs using multiple controllers types)")
    else:
        add_p().add_run("There are no VMs with multiple controllers types")

    return


def print_storage_hba_models(sheets_dict):
    ##################################
    # HBA Models
    #TODO
    add_p().add_run("")
    rowLink = []
    dataList = calculate_percentage(sheets_dict, sheet='vHBA', columns=['Model', 'Type'])
    for value, count, percentage in dataList:
        rowLink.append("https://www.google.com/search?q==" + urllib.parse.quote(value[0], safe='/', encoding=None,
                                                                                      errors=None))
    add_percentage_table(document, dataList, "[vHBA] Model/Type", None, "Table Grid", None, rowLink)
    return


def print_storage_checks_and_hints(sheets_dict, compute_sizing_totals_row):
    # Storage Checks and Hints
    # #####################################################################################################################
    storage_checks = {}
    add_p().add_run("")

    total_number_of_Disks = str(count_rows(sheets_dict, 'vDisk'))
    total_number_of_Partitions = str(count_rows(sheets_dict, 'vPartition'))

    key = "VMs with more than 20 disks (vInfo)"
    try:
        storage_checks[key] = count_rows(sheets_dict, 'vInfo', query_expr='Disks > 20')
    except Exception as ex:
        storage_checks[key] = str(ex)

    key = "Disks between 1TiB and 2TiB (vDisk)"
    try:
        storage_checks[key] = count_rows(sheets_dict, 'vDisk',
                                         query_expr='Capacity_MiB > 1048576 and Capacity_MiB <= 2097152')
    except Exception as ex:
        storage_checks[key] = str(ex)

    key = "Disks between 2TB and 4TB (vDisk)"
    try:
        storage_checks[key] = count_rows(sheets_dict, 'vDisk',
                                         query_expr='Capacity_MiB > 2097152 and Capacity_MiB <= 4194304')
    except Exception as ex:
        storage_checks[key] = str(ex)

    key = "Disks between 4TB and 6TB (vDisk)"
    try:
        storage_checks[key] = count_rows(sheets_dict, 'vDisk',
                                         query_expr='Capacity_MiB > 4194304 and Capacity_MiB <= 6291456')
    except Exception as ex:
        storage_checks[key] = str(ex)

    key = "Disks with more than 6TB (vDisk)"
    try:
        large_disks = None
        storage_checks[key] = count_rows(sheets_dict, 'vDisk', query_expr='Capacity_MiB > 6291456')
        if storage_checks[key] > 0:
            large_disks = get_rows(sheets_dict, 'vPartition', key_columns=['Annotation', 'Capacity_MiB', 'Disk'],
                                   columns=['VM', 'Disk', 'Capacity_MiB', 'Annotation'],
                                   query_expr='Capacity_MiB > 6291456')
            large_disks.sort_values(by=['Capacity_MiB', 'Disk', 'Annotation'], ascending=False, inplace=True)
            # print(large_disks.head(50))

    except Exception as ex:
        storage_checks[key] = str(ex)

    key = "Partitions with less than 10 Mib free (vPartition)"
    try:
        storage_checks[key] = count_rows(sheets_dict, sheet='vPartition', query_expr="Free_MiB < 10")
    except Exception as ex:
        storage_checks[key] = str(ex)

    key = "Partitions with less than 50 Mib free (vPartition)"
    try:
        storage_checks[key] = count_rows(sheets_dict, sheet='vPartition', query_expr="Free_MiB < 50")
    except Exception as ex:
        storage_checks[key] = str(ex)

    key = "Partitions with less than 100 Mib free (vPartition)"
    try:
        storage_checks[key] = count_rows(sheets_dict, sheet='vPartition', query_expr="Free_MiB < 100")
    except Exception as ex:
        storage_checks[key] = str(ex)

    key = "Root Partitions with less than 100 Mib free (vPartition)"
    try:
        storage_checks[key] = count_rows(sheets_dict, sheet='vPartition',
                                         query_expr="(Disk == '/' or Disk == 'C:\\\\') and Free_MiB < 100")
    except Exception as ex:
        storage_checks[key] = str(ex)

    key = "Root Partitions with less than 200 Mib free (vPartition)"
    try:
        storage_checks[key] = count_rows(sheets_dict, sheet='vPartition',
                                         query_expr="(Disk == '/' or Disk == 'C:\\\\') and Free_MiB < 200")
    except Exception as ex:
        storage_checks[key] = str(ex)

    key = "Root Partitions with less than 500 Mib free (vPartition)"
    try:
        storage_checks[key] = count_rows(sheets_dict, sheet='vPartition',
                                         query_expr="(Disk == '/' or Disk == 'C:\\\\') and Free_MiB < 500")
    except Exception as ex:
        storage_checks[key] = str(ex)

    table = table_from_dict(storage_checks, document, True, None,
                            "Check", "Data", True, "Storage Checks & Hints")

    # TODO: Avoid this ugly and crazy way to reuse information
    total_number_of_VMs = compute_sizing_totals_row.cells[9].text

    run = add_p().add_run("Reference:  Total VMs=" + total_number_of_VMs +
                         "   Total Disks=" + total_number_of_Disks +
                         "   Total Partitions=" + total_number_of_Partitions)
    run.font.color.rgb = gray_color

    return


def print_storage_large_disks(sheets_dict):
    ################################################################################
    # TOP Large Disks List
    ################################################################################

    # Retrieve information
    size = 4000000
    large_disks = get_rows(sheets_dict, 'vPartition', key_columns=None,
                           columns=['VM', 'Disk', 'Capacity_MiB', 'Annotation'],
                           query_expr='Capacity_MiB > ' + str(size))

    if large_disks.size > 0:
        large_disks.sort_values(by=['Capacity_MiB', 'Disk', 'Annotation'], ascending=False, inplace=True)

        ldisks = []
        pre_vm = ''
        pre_disk = ''
        pre_capacity = ''
        pre_annotation = ''
        vm = ''
        count = 1

        # Add headers
        ldisks.append(["Disk", "GiB", "VMs", "Annotations", "Cnt"])

        # Add content
        for index, row in large_disks.iterrows():
            disk = row['Disk']
            capacity = f'{int(int(row["Capacity_MiB"]) / 1024):n}'
            vm = row['VM']
            annotation = row['Annotation']

            if pre_capacity != capacity or pre_disk != disk:
                if pre_disk != '':  # do not add the first empty element
                    ldisks.append([pre_disk, pre_capacity, pre_vm, pre_annotation, count])
                pre_vm = ''
                count = 1
            else:
                count += 1

            pre_disk = disk
            pre_capacity = capacity
            pre_vm = pre_vm + vm + ' '
            if len(str(annotation)) > 5: pre_annotation = annotation

        ldisks.append([pre_disk, pre_capacity, pre_vm, pre_annotation, count])

        # Write to DOC
        if large_disks is not None:
            top = 40
            add_p().add_run("")

            L = WD_ALIGN_PARAGRAPH.LEFT
            R = WD_ALIGN_PARAGRAPH.RIGHT
            table = table_from_list(ldisks[:top], document, True, [100, 35, 170, 170, 25], [L, R, L, L, R],
                                    "[vPartition] Top " + str(
                                        top) + " large disks greater than " + f'{int(int(size / 1000)):n}' + " GiB")

            # Adjust fonts and lengths
            # 0:Disk - 1:Capacity - 2:VMs - 3:Annotations - 4:Count
            count = 0
            max_len = 130
            for row in table.rows:
                count += 1
                if count > 2:  # jump headers
                    for cell in row.cells: cell.paragraphs[0].runs[0].font.size = Pt(8)

                    if len(str(row.cells[0].text)) > 22: row.cells[0].paragraphs[0].runs[0].font.size = Pt(7)
                    if len(str(row.cells[0].text)) > 26: row.cells[0].paragraphs[0].runs[0].font.size = Pt(6)

                    if len(str(row.cells[2].text)) > max_len: row.cells[2].text = row.cells[2].text[0:max_len] + "..."
                    if len(str(row.cells[2].text)) > 70: row.cells[2].paragraphs[0].runs[0].font.size = Pt(7)
                    if len(str(row.cells[2].text)) > 120: row.cells[2].paragraphs[0].runs[0].font.size = Pt(6)

                    if len(str(row.cells[3].text)) > max_len: row.cells[3].text = row.cells[3].text[0:max_len] + "..."
                    if len(str(row.cells[3].text)) > 70: row.cells[3].paragraphs[0].runs[0].font.size = Pt(7)
                    if len(str(row.cells[3].text)) > 100: row.cells[3].paragraphs[0].runs[0].font.size = Pt(6)
    else:
        add_p().add_run("There are no Disks greater than " + f'{int(int(size / 1000)):n}' + " GiB")

    return


def print_anonymize_data(document, data):
    """
    Print anonymize information to revert it
    :param document:
    :param data:
    :return:
    """
    ##########################################################
    # %%% PRINT Anonymization information
    if data is not None:
        document.add_page_break()
        document.add_heading("DataCenter/Cluster Masking Table", 1)
        # mylist=[["Original","Mask"]]
        mylist = []

        anonym_datacenter = data.get("datacenter")
        # print(anonym_datacenter)
        for key in anonym_datacenter.keys():
            row = [key, anonym_datacenter[key]]
            mylist.append(row)

        anonym_cluster = data.get("cluster")
        # print(anonym_cluster)
        for key in anonym_cluster.keys():
            row = [key, anonym_cluster[key]]
            mylist.append(row)

        # print(mylist)
        table = table_from_list(mylist, document, True, None, None)

        return

def build_docx(sheets_dict, anonymize_data, output_file=None):
    """
    Fill the global document from an already-loaded and cleaned sheets_dict.
    If output_file is set, save there and return None; otherwise return BytesIO.
    """
    global document, warns

    warns.clear()
    document = Document()

    print_initialize(document)
    lapse("print_initialize()")

    print_versions(sheets_dict)
    lapse("print_versions()")

    print_vmw_products(sheets_dict)
    lapse("print_vmw_products()")

    compute_sizing_totals_row = print_compute_sizing(sheets_dict)
    lapse("compute_sizing_totals_row()")

    print_memory_ranks(sheets_dict)
    lapse("print_memory_ranks()")

    print_compute_checks_and_hints(sheets_dict, compute_sizing_totals_row)
    lapse("print_compute_checks_and_hints()")

    print_networking(sheets_dict)
    lapse("print_networking()")

    print_network_terms(sheets_dict)
    lapse("print_network_terms()")

    print_workload_terms(sheets_dict)
    lapse("print_workload_terms()")

    print_appliances_n_OVA_annotations(sheets_dict)
    lapse("print_appliances_n_OVA_annotations()")

    print_operating_systems(sheets_dict)
    lapse("print_operating_systems()")

    print_storage_models(sheets_dict)
    lapse("print_storage_models()")

    print_storage_capacity(sheets_dict)
    lapse("print_storage_capacity()")

    print_storage_connections(sheets_dict)
    lapse("print_storage_connections()")

    print_storage_vms_per_disk_controller(sheets_dict)
    lapse("print_storage_vms_per_disk_controller()")

    print_storage_multiple_controllers(sheets_dict)
    lapse("print_storage_multiple_controllers()")

    print_storage_hba_models(sheets_dict)
    lapse("print_storage_hba_models()")

    print_storage_checks_and_hints(sheets_dict, compute_sizing_totals_row)
    lapse("print_storage_checks_and_hints()")

    print_storage_large_disks(sheets_dict)
    lapse("print_storage_large_disks()")

    if len(warns) > 0:
        document.add_page_break()
        document.add_heading("Warnings during report processing", 1)
        p = add_p()
        for warn in warns:
            p.add_run(warn + "\n")

    print_anonymize_data(document, anonymize_data)

    keep_table_on_one_page(document)

    if output_file is not None:
        document.save(output_file)
        print("File saved: " + output_file)
        return None

    buf = BytesIO()
    document.save(buf)
    buf.seek(0)
    return buf


def run_pipeline_from_bytes(content, original_filename, anonymize=False, lang="en"):
    """
    End-to-end: bytes -> validated sheets -> clean -> docx in memory.
    lang: 'en' or 'de' for upload/validation error messages.
    """
    pd.options.display.max_rows = 1000
    pd.options.display.max_columns = 20
    pd.options.display.width = 2000
    pd.set_option("display.expand_frame_repr", True)

    sheets_dict = load_spreadsheets_bytes(content, original_filename, lang=lang)
    clean_and_fix_data(sheets_dict)
    anonymize_data = anonymize_names(sheets_dict, anonymize)
    return build_docx(sheets_dict, anonymize_data, output_file=None)


def main():
    import os

    # Get the arguments from the command-line except the filename
    argv = sys.argv
    path = None
    print("AnalytiX version:" + str(version))
    if len(sys.argv) < 2:
        print("Parameters:")
        print(
            "   Param1: Path_to_a_folder containing the *.xlsx and/or *.xlsm files to be processed. Use '.' for current folder")
        print('   Param2 (Optional): name of the resulting file. "analytx.docx" by default')
        # argv.append("/home/marmendo/Documentos/PROYECTOS/FeasibilityReport/Customers/ACME/RVTools")
        exit()

    path = argv[1]

    if len(sys.argv) > 2:
        output_file = argv[2]
    else:
        output_file = "./analytx.docx"

    lapse(on=False)

    pd.options.display.max_rows = 1000
    pd.options.display.max_columns = 20
    pd.options.display.width = 2000
    pd.set_option("display.expand_frame_repr", True)
    anonymize = False  # TODO: anonymize doesn't work as expected... yet

    sheets_dict = load_spreadsheets(path)
    lapse("load_spreadsheets()")
    clean_and_fix_data(sheets_dict)
    lapse("clean_and_fix_data()")
    anonymize_data = anonymize_names(sheets_dict, anonymize)
    lapse("anonymize_names()")

    build_docx(sheets_dict, anonymize_data, output_file=output_file)

    # Don't delete this line. See cols_prepare
    # print(cols_prepare.idx_sheets)

    return


if __name__ == "__main__":
    main()
